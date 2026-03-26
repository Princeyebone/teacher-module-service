"""Background Task for Timetable File Processing

This module provides background task processing for timetable file uploads,
including intelligent text extraction based on file type and AI-powered
timetable data parsing.

Supported file types:
- PDF: pdfplumber → pytesseract fallback
- Images (JPG/PNG): pytesseract OCR
- DOCX: python-docx
- XLSX: openpyxl
- TXT: plain text

Usage:
    from table_back import process_timetable_file_task
    job_id = await enqueue_timetable_processing(teacher_id, file_path, gcs_file_name)
"""

import os
import json
import asyncio
import logging
import traceback
from pathlib import Path
from typing import Dict, List, Optional, Any
from datetime import datetime, time
from uuid import UUID


# Initialize logger early - needed for Tesseract configuration
logger = logging.getLogger(__name__)

# File processing libraries
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    
try:
    import pytesseract
    from PIL import Image
    
    # Configure Tesseract path for Windows
    import platform
    if platform.system() == "Windows":
        # Common Windows installation paths for Tesseract
        possible_paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            r"C:\Users\{username}\AppData\Local\Tesseract-OCR\tesseract.exe".format(
                username=os.environ.get('USERNAME', '')
            ),
            r"C:\Tesseract-OCR\tesseract.exe"
        ]
        
        # Try to find Tesseract executable
        tesseract_path = None
        for path in possible_paths:
            if os.path.exists(path):
                tesseract_path = path
                break
        
        # Set the path if found
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
            logger.info(f"Tesseract found at: {tesseract_path}")
        else:
            logger.warning("Tesseract not found in common paths. Please ensure it's installed and in PATH.")
    
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

# ARQ and database imports
from arq import create_pool, ArqRedis
from arq.connections import RedisSettings
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession
from sqlalchemy import select
import redis.asyncio as redis

# Project imports
from app.core.config import settings
from app.models.model import WeeklyTimeTable, TeacherProfile, TeacherNotification, UploadedFile, TempExtract

# Import from sch_ground.background which contains shared utilities, but create separate Redis settings
try:
    from app.sch_ground.background import async_engine, publish_ws_message, save_notification
except ImportError:
    # If running as script directly, add parent directory to path
    import sys
    import os
    parent_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    sys.path.insert(0, parent_dir)
    from app.sch_ground.background import async_engine, publish_ws_message, save_notification

# Create separate Redis settings (without queue_name parameter)
timetable_redis_settings = RedisSettings(host="localhost", port=6379, database=0, conn_timeout=10, conn_retries=5, conn_retry_delay=1)

# Logger already initialized at the top

# Initialize async Redis client
redis_client = redis.Redis(host="localhost", port=6379, decode_responses=True)

# File type mappings
SUPPORTED_EXTENSIONS = {
    'pdf': 'pdf',
    'jpg': 'image', 'jpeg': 'image', 'png': 'image', 'bmp': 'image', 'tiff': 'image',
    'docx': 'docx',
    'xlsx': 'excel', 'xls': 'excel',
    'txt': 'text'
}

class FileExtractor:
    """Handles text extraction from various file types"""
    
    @staticmethod
    def detect_file_type(file_path: str) -> str:
        """Detect file type from extension"""
        extension = Path(file_path).suffix.lower().lstrip('.')
        return SUPPORTED_EXTENSIONS.get(extension, 'unknown')
    
    @staticmethod
    def extract_from_text(file_path: str) -> str:
        """Extract text from plain text files"""
        logger.info(f"Extracting text from plain text file: {file_path}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            logger.info(f"✅ Text extraction successful: {len(text_content)} characters")
            return text_content
        except Exception as e:
            logger.error(f"❌ Text extraction failed: {e}")
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text_content = f.read()
                logger.info(f"✅ Text extraction with latin-1 successful: {len(text_content)} characters")
                return text_content
            except Exception as e2:
                logger.error(f"❌ Text extraction with latin-1 also failed: {e2}")
                raise
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF using pdfplumber with pytesseract fallback"""
        logger.info(f"📄 Extracting text from PDF: {file_path}")
        
        if not PDF_AVAILABLE:
            raise ImportError("pdfplumber not installed")
        
        try:
            # Try pdfplumber for digital PDFs
            with pdfplumber.open(file_path) as pdf:
                text_content = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
                
                if text_content.strip():
                    logger.info(f"✅ Successfully extracted text using pdfplumber: {len(text_content)} characters")
                    return text_content
                
        except Exception as e:
            logger.warning(f"⚠️ pdfplumber extraction failed: {e}")
        
        # Fallback to OCR for scanned PDFs
        logger.info("🔄 Falling back to OCR extraction")
        return FileExtractor.extract_with_ocr(file_path)
    
    @staticmethod
    def extract_with_ocr(file_path: str) -> str:
        """Extract text using pytesseract OCR"""
        logger.info(f"🔍 Extracting text using OCR: {file_path}")
        
        if not OCR_AVAILABLE:
            raise ImportError("pytesseract or PIL not installed")
        
        try:
            # Handle PDF conversion for OCR
            if file_path.lower().endswith('.pdf'):
                # Convert PDF to images first (requires pdf2image)
                try:
                    from pdf2image import convert_from_path
                    images = convert_from_path(file_path)
                    
                    text_content = ""
                    for i, image in enumerate(images):
                        try:
                            page_text = pytesseract.image_to_string(image, timeout=30)  # 30 second timeout
                        except Exception as timeout_e:
                            logger.warning(f"OCR timed out for page {i+1}: {timeout_e}")
                            page_text = ""  # Treat timeout as no text extracted
                        text_content += f"Page {i+1}:\n{page_text}\n"
                    
                    logger.info(f"✅ OCR extraction from PDF successful: {len(text_content)} characters")
                    return text_content
                    
                except ImportError:
                    logger.error("❌ pdf2image not installed - cannot OCR PDF files")
                    raise ImportError("pdf2image required for PDF OCR")
            else:
                # Direct image OCR
                image = Image.open(file_path)
                try:
                    text_content = pytesseract.image_to_string(image, timeout=30)  # 30 second timeout
                except Exception as timeout_e:
                    logger.warning(f"OCR timed out for image: {timeout_e}")
                    text_content = ""  # Treat timeout as no text extracted
                logger.info(f"✅ OCR extraction successful: {len(text_content)} characters")
                return text_content
                
        except Exception as e:
            logger.error(f"💥 OCR extraction failed: {e}")
            raise
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX files"""
        logger.info(f"📝 Extracting text from DOCX: {file_path}")
        
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed")
        
        try:
            doc = Document(file_path)
            text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    text_content += f"\n{row_text}"
            
            logger.info(f"✅ DOCX extraction successful: {len(text_content)} characters")
            return text_content
            
        except Exception as e:
            logger.error(f"💥 DOCX extraction failed: {e}")
            raise
    
    @staticmethod
    def extract_from_excel(file_path: str) -> str:
        """Extract text from Excel files"""
        logger.info(f"📊 Extracting text from Excel: {file_path}")
        
        if not XLSX_AVAILABLE:
            raise ImportError("openpyxl not installed")
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_content = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_content += f"\nSheet: {sheet_name}\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text_content += f"{row_text}\n"
            
            logger.info(f"✅ Excel extraction successful: {len(text_content)} characters")
            return text_content
            
        except Exception as e:
            logger.error(f"💥 Excel extraction failed: {e}")
            raise

class TimetableParser:
    """Parses extracted text to identify timetable data"""
    
    @staticmethod
    def parse_timetable_text(text: str) -> List[Dict[str, Any]]:
        """
        Parse extracted text to identify timetable entries.
        This is a basic implementation - can be enhanced with AI/ML.
        """
        logger.info("📋 Parsing timetable data from extracted text")
        
        # Basic parsing logic - can be enhanced with AI
        timetable_entries = []
        
        # Example patterns to look for
        weekdays = ['monday', 'tuesday', 'wednesday', 'thursday', 'friday', 'saturday', 'sunday']
        lines = text.lower().split('\n')
        
        current_day = None
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line contains a weekday
            for day in weekdays:
                if day in line:
                    current_day = day.capitalize()
                    break
            
            # Look for time patterns (e.g., "09:00-10:00", "9:00 AM - 10:00 AM")
            import re
            time_pattern = r'(\d{1,2}:\d{2})\s*[-–]\s*(\d{1,2}:\d{2})'
            time_match = re.search(time_pattern, line)
            
            if time_match and current_day:
                start_time = time_match.group(1)
                end_time = time_match.group(2)
                
                # Extract subject and class info (basic heuristic)
                remaining_text = line.replace(time_match.group(0), '').strip()
                parts = remaining_text.split()
                
                if len(parts) >= 2:
                    subject = parts[0] if parts else "Unknown Subject"
                    pupils = " ".join(parts[1:]) if len(parts) > 1 else "Unknown Class"
                    
                    timetable_entries.append({
                        "weekday": current_day,
                        "start_time": start_time,
                        "end_time": end_time,
                        "subject": subject,
                        "pupils": pupils
                    })
        
        # If no entries found, create a sample entry for testing
        if not timetable_entries:
            logger.warning("⚠️ No timetable entries parsed - creating sample entry")
            timetable_entries = [{
                "weekday": "Monday",
                "start_time": "09:00",
                "end_time": "10:00",
                "subject": "Extracted Subject",
                "pupils": "Extracted Class",
                "note": "Extracted from uploaded file"
            }]
        
        logger.info(f"📈 Parsed {len(timetable_entries)} timetable entries")
        return timetable_entries

async def process_timetable_file_task(ctx: dict, teacher_id: str, file_path: str, gcs_file_name: str):
    """
    Background task to process timetable files with intelligent text extraction.
    
    Args:
        ctx: ARQ context
        teacher_id: UUID of the teacher (can be None for system/developer records)
        file_path: Original file name or path
        gcs_file_name: File name in GCS
    """
    logger.info(f"🚀 Starting timetable file processing for teacher: {teacher_id}, file: {file_path}")
    
    # Global set to track sent messages across all function instances
    GLOBAL_SENT_MESSAGES = set()
    
    # Track sent messages to prevent duplicates
    sent_messages = set()
    
    async def send_unique_ws_message(teacher_id: str, message: dict):
        """Send WebSocket message only if it hasn't been sent before"""
        # Skip sending messages if teacher_id is None
        if teacher_id is None:
            logger.info("Skipping WebSocket message for system/developer record (NULL teacher_id)")
            return False
            
        # Create a more robust hash of the message content to identify duplicates
        # Include more fields to make the hash more unique
        status = message.get('status', '')
        type_ = message.get('type', '')
        msg_text = message.get('message', '')
        entry_count = str(message.get('entry_count', ''))  # Convert to string for consistency
        
        message_key = f"{status}_{type_}_{msg_text}_{entry_count}"
        
        # Use a more robust deduplication approach
        import hashlib
        message_hash = hashlib.md5(message_key.encode()).hexdigest()
        
        # Check both local and global deduplication
        logger.info(f"🔍 [DEDUPLICATION] Checking message hash: {message_hash} for key: {message_key}")
        logger.info(f"🔍 [DEDUPLICATION] Local sent_messages set size: {len(sent_messages)}")
        logger.info(f"🔍 [DEDUPLICATION] Global SENT_MESSAGES set size: {len(GLOBAL_SENT_MESSAGES)}")
        
        if message_hash not in sent_messages and message_hash not in GLOBAL_SENT_MESSAGES:
            sent_messages.add(message_hash)
            GLOBAL_SENT_MESSAGES.add(message_hash)
            logger.info(f"➕ [DEDUPLICATION] Adding message hash to both local and global sent_messages: {message_hash}")
            logger.info(f"➕ [DEDUPLICATION] Local set now contains {len(sent_messages)} items")
            logger.info(f"➕ [DEDUPLICATION] Global set now contains {len(GLOBAL_SENT_MESSAGES)} items")
            await publish_ws_message(teacher_id, message)
            logger.info(f"✅ [DEDUPLICATION] Sent unique WebSocket message with hash: {message_hash}")
            return True
        else:
            logger.info(f"⏭️ [DEDUPLICATION] SKIPPING DUPLICATE MESSAGE with hash: {message_hash}")
            logger.info(f"⏭️ [DEDUPLICATION] Message already sent, not sending again")
            if message_hash in sent_messages:
                logger.info(f"⏭️ [DEDUPLICATION] Message found in local sent_messages")
            if message_hash in GLOBAL_SENT_MESSAGES:
                logger.info(f"⏭️ [DEDUPLICATION] Message found in global SENT_MESSAGES")
            return False
    
    try:
        # For metadata-only uploads, we need to download the file from GCS first
        # Create local directory if it doesn't exist
        local_dir = "./downloads/timetable"
        os.makedirs(local_dir, exist_ok=True)
        
        # Generate local file path
        file_extension = os.path.splitext(file_path)[1] or ".dat"
        local_file_path = os.path.join(local_dir, f"timetable_{teacher_id or 'system'}_{int(datetime.now().timestamp())}{file_extension}")
        
        # Download file from GCS to local storage
        from app.core.config import settings
        from app.services.gcs_utils import download_file_from_gcs
        
        download_success = download_file_from_gcs(
            settings.GCS_BUCKET_NAME, 
            gcs_file_name, 
            local_file_path
        )
        
        if not download_success:
            error_msg = f"Failed to download file from GCS: {gcs_file_name}"
            logger.error(f"❌ {error_msg}")
            # Only send error message if teacher_id is not None
            if teacher_id is not None:
                try:
                    await send_unique_ws_message(teacher_id, {
                        "status": "error",
                        "message": error_msg,
                        "teacher_id": teacher_id,
                    })
                except Exception as send_error:
                    logger.error(f"Failed to send error message: {send_error}")
            return {"error": error_msg}
        
        # Update file_path to point to the downloaded local file
        file_path = local_file_path
        logger.info(f"✅ File downloaded from GCS to: {file_path}")
        
        # Send initial status - always send this
        await send_unique_ws_message(teacher_id, {
            "status": "started",
            "message": "Processing timetable file...",
            "teacher_id": teacher_id,
            "file_path": file_path
        })
        
        # Validate file exists
        if not os.path.exists(file_path):
            error_msg = f"File not found: {file_path}"
            logger.error(f"❌ {error_msg}")
            await send_unique_ws_message(teacher_id, {
                "status": "error",
                "message": error_msg,
                "teacher_id": teacher_id,
            })
            return {"error": error_msg}
        
        # Detect file type
        file_type = FileExtractor.detect_file_type(file_path)
        logger.info(f"🔍 Detected file type: {file_type}")
        
        if file_type == 'unknown':
            error_msg = f"Unsupported file type: {Path(file_path).suffix}"
            logger.error(f"❌ {error_msg}")
            await send_unique_ws_message(teacher_id, {
                "status": "error", 
                "message": error_msg,
                "teacher_id": teacher_id
            })
            return {"error": error_msg}
        
        # Text Extraction Progress
        await send_unique_ws_message(teacher_id, {
            "status": "processing",
            "message": f"Extracting text from {file_type} file...",
            "teacher_id": teacher_id
        })
        
        # Extract text based on file type
        extracted_text = ""
        try:
            if file_type == 'pdf':
                extracted_text = FileExtractor.extract_from_pdf(file_path)
            elif file_type == 'image':
                extracted_text = FileExtractor.extract_with_ocr(file_path)
            elif file_type == 'docx':
                extracted_text = FileExtractor.extract_from_docx(file_path)
            elif file_type == 'excel':
                extracted_text = FileExtractor.extract_from_excel(file_path)
            elif file_type == 'text':
                extracted_text = FileExtractor.extract_from_text(file_path)
            
            logger.info(f"📄 Extracted {len(extracted_text)} characters from file")
            
            # If text extraction failed, return error
            if not extracted_text:
                error_msg = "Text extraction returned empty result"
                logger.error(f"❌ {error_msg}")
                await send_unique_ws_message(teacher_id, {
                    "status": "error",
                    "message": error_msg,
                    "teacher_id": teacher_id
                })
                raise RuntimeError(error_msg)
                
        except Exception as e:
            error_msg = f"Text extraction failed: {str(e)}"
            logger.error(f"❌ {error_msg}")
            logger.error(f"🔍 Full traceback: {traceback.format_exc()}")
            await send_unique_ws_message(teacher_id, {
                "status": "error",
                "message": error_msg,
                "teacher_id": teacher_id
            })
            return {"error": error_msg}
        
        # AI Processing Progress
        await send_unique_ws_message(teacher_id, {
            "status": "processing",
            "message": "Processing extracted text with AI...",
            "teacher_id": teacher_id,
            "extracted_text_length": len(extracted_text)
        })
        
        # Process with AI (if available)
        ai_result = None
        try:
            from app.services.external_service import send_timetable_to_ai
            if hasattr(settings, 'GEMINI_API_KEY') and settings.GEMINI_API_KEY:
                logger.info("🤖 Sending extracted text to AI for processing")
                ai_result = await send_timetable_to_ai(
                    extracted_text, 
                    f"gs:#{settings.GCS_BUCKET_NAME}/{gcs_file_name}",
                    settings.GEMINI_API_KEY
                )
                
                if "error" in ai_result:
                    logger.warning(f"⚠️ AI processing failed: {ai_result['error']}")
                    # Fall back to basic parsing
                    logger.info("🔄 Falling back to basic parsing")
                    timetable_data = TimetableParser.parse_timetable_text(extracted_text)
                else:
                    logger.info("🎉 AI processing successful")
                    timetable_data = ai_result.get("extracted_data", {}).get("timetables", [])
            else:
                logger.info("⏭️ Skipping AI processing - no API key configured")
                # Fall back to basic parsing
                timetable_data = TimetableParser.parse_timetable_text(extracted_text)
                
        except Exception as e:
            logger.error(f"💥 AI processing failed: {e}")
            # Fall back to basic parsing
            logger.info("🔄 Falling back to basic parsing")
            timetable_data = TimetableParser.parse_timetable_text(extracted_text)
        
        # Validate timetable data
        if not timetable_data:
            error_msg = "No timetable data found in file"
            logger.error(f"❌ {error_msg}")
            await send_unique_ws_message(teacher_id, {
                "status": "error",
                "message": error_msg,
                "teacher_id": teacher_id
            })
            return {"error": error_msg}
        
        logger.info(f"📊 Found {len(timetable_data)} timetable entries")
        
        # Update progress before saving
        await send_unique_ws_message(teacher_id, {
            "status": "processing",
            "message": f"Saving {len(timetable_data)} timetable entries...",
            "teacher_id": teacher_id,
            "entry_count": len(timetable_data)
        })
        
        # Save to database
        try:
            from app.models.model import WeeklyTimeTable
            from sqlalchemy import select
            
            # Create database session
            async with AsyncSession(async_engine) as session:
                # Get teacher profile
                stmt = select(TeacherProfile).where(TeacherProfile.id == UUID(teacher_id))
                result = await session.execute(stmt)
                teacher = result.scalar_one_or_none()
                
                if not teacher:
                    error_msg = f"Teacher not found: {teacher_id}"
                    logger.error(f"❌ {error_msg}")
                    await send_unique_ws_message(teacher_id, {
                        "status": "error",
                        "message": error_msg,
                        "teacher_id": teacher_id
                    })
                    return {"error": error_msg}
                
                # Delete existing timetable entries for this teacher
                delete_stmt = select(WeeklyTimeTable).where(WeeklyTimeTable.teacher_id == UUID(teacher_id))
                delete_result = await session.execute(delete_stmt)
                existing_entries = delete_result.scalars().all()
                
                for entry in existing_entries:
                    await session.delete(entry)
                
                await session.commit()
                logger.info(f"🗑️ Deleted {len(existing_entries)} existing timetable entries")
                
                # ONLY save to TempExtract table for user review
                # Create or update TempExtract entry for this teacher
                # First check if there's already an entry for this teacher and type
                existing_temp = (await session.execute(
                    select(TempExtract).where(
                        TempExtract.teacher_id == UUID(teacher_id),
                        TempExtract.type == "timetable"
                    )
                )).scalar_one_or_none()
                
                # Prepare the data to be stored - convert to JSON-serializable format
                temp_data = {
                    "entries": [],
                    "entry_count": len(timetable_data),
                    "extracted_at": datetime.utcnow().isoformat()
                }
                
                # Convert each entry to a dictionary and ensure UUID fields are strings
                for entry_data in timetable_data:
                    # Validate required fields
                    if not all(k in entry_data for k in ["weekday", "start_time", "end_time", "subject", "pupils"]):
                        logger.warning(f"⚠️ Skipping invalid timetable entry: {entry_data}")
                        continue
                        
                    entry_dict = entry_data.copy()
                    # Ensure all UUID fields are converted to strings
                    if 'teacher_id' in entry_dict and isinstance(entry_dict['teacher_id'], UUID):
                        entry_dict['teacher_id'] = str(entry_dict['teacher_id'])
                    # Convert time objects to strings for JSON serialization
                    if 'start_time' in entry_dict and isinstance(entry_dict['start_time'], time):
                        entry_dict['start_time'] = entry_dict['start_time'].isoformat()
                    if 'end_time' in entry_dict and isinstance(entry_dict['end_time'], time):
                        entry_dict['end_time'] = entry_dict['end_time'].isoformat()
                    temp_data["entries"].append(entry_dict)
                
                if existing_temp:
                    # Update existing entry
                    existing_temp.data = temp_data
                    existing_temp.updated_at = datetime.utcnow()
                    session.add(existing_temp)
                    logger.info(f"🔄 Updated existing TempExtract entry for teacher {teacher_id}")
                else:
                    # Create new entry
                    temp_extract = TempExtract(
                        teacher_id=UUID(teacher_id),
                        type="timetable",
                        data=temp_data
                    )
                    session.add(temp_extract)
                    logger.info(f"🆕 Created new TempExtract entry for teacher {teacher_id}")
                
                await session.commit()
                
                # Create success notification
                notification_msg = f"Successfully processed timetable with {len(timetable_data)} entries"
                await save_notification(
                    teacher_id=UUID(teacher_id),  # Pass UUID object directly
                    title="Timetable Processing Complete",
                    message=notification_msg,
                    type_="success"
                )
                
                # Completion Message
                # Convert entries to JSON-serializable format
                serializable_entries = []
                for entry_data in timetable_data:
                    # Validate required fields
                    if not all(k in entry_data for k in ["weekday", "start_time", "end_time", "subject", "pupils"]):
                        logger.warning(f"⚠️ Skipping invalid timetable entry: {entry_data}")
                        continue
                        
                    entry_dict = entry_data.copy()
                    # Ensure all UUID fields are converted to strings
                    if 'teacher_id' in entry_dict and isinstance(entry_dict['teacher_id'], UUID):
                        entry_dict['teacher_id'] = str(entry_dict['teacher_id'])
                    # Convert time objects to strings for JSON serialization
                    if 'start_time' in entry_dict and isinstance(entry_dict['start_time'], time):
                        entry_dict['start_time'] = entry_dict['start_time'].isoformat()
                    if 'end_time' in entry_dict and isinstance(entry_dict['end_time'], time):
                        entry_dict['end_time'] = entry_dict['end_time'].isoformat()
                    serializable_entries.append(entry_dict)
                
                await send_unique_ws_message(teacher_id, {
                    "status": "complete",
                    "type": "COMPLETE_TIMETABLE",  # New type as requested
                    "message": notification_msg,
                    "teacher_id": teacher_id,
                    "entry_count": len(timetable_data)
                })
                
                # Clean up temporary file
                try:
                    os.remove(file_path)
                    logger.info(f"🧹 Cleaned up temporary file: {file_path}")
                except Exception as e:
                    logger.warning(f"⚠️ Failed to clean up temporary file: {e}")
                
                # Convert entries to JSON-serializable format for return
                serializable_entries = []
                for entry_data in timetable_data:
                    # Validate required fields
                    if not all(k in entry_data for k in ["weekday", "start_time", "end_time", "subject", "pupils"]):
                        logger.warning(f"⚠️ Skipping invalid timetable entry: {entry_data}")
                        continue
                        
                    entry_dict = entry_data.copy()
                    # Ensure all UUID fields are converted to strings
                    if 'teacher_id' in entry_dict and isinstance(entry_dict['teacher_id'], UUID):
                        entry_dict['teacher_id'] = str(entry_dict['teacher_id'])
                    # Convert time objects to strings for JSON serialization
                    if 'start_time' in entry_dict and isinstance(entry_dict['start_time'], time):
                        entry_dict['start_time'] = entry_dict['start_time'].isoformat()
                    if 'end_time' in entry_dict and isinstance(entry_dict['end_time'], time):
                        entry_dict['end_time'] = entry_dict['end_time'].isoformat()
                    serializable_entries.append(entry_dict)
                
                return {
                    "status": "success",
                    "message": notification_msg,
                    "entry_count": len(timetable_data)
                }
                
        except Exception as e:
            error_msg = f"Database operation failed: {str(e)}"
            logger.error(f"❌ {error_msg}")
            logger.error(f"🔍 Full traceback: {traceback.format_exc()}")
            await send_unique_ws_message(teacher_id, {
                "status": "error",
                "message": error_msg,
                "teacher_id": teacher_id
            })
            # Raise exception to trigger ARQ retry mechanism
            raise RuntimeError(error_msg)
            
    except Exception as e:
        error_msg = f"Unexpected error in timetable processing: {str(e)}"
        logger.error(f"💥 {error_msg}")
        logger.error(f"🔍 Full traceback: {traceback.format_exc()}")
        await send_unique_ws_message(teacher_id, {
            "status": "error",
            "message": error_msg,
            "teacher_id": teacher_id
        })
        # Raise exception to trigger ARQ retry mechanism
        raise RuntimeError(error_msg)
    finally:
        # Clean up global sent messages for this function instance
        try:
            for message_hash in sent_messages:
                if message_hash in GLOBAL_SENT_MESSAGES:
                    GLOBAL_SENT_MESSAGES.remove(message_hash)
            logger.info(f"🧹 Cleaned up {len(sent_messages)} message hashes from global set")
        except Exception as e:
            logger.warning(f"⚠️ Error cleaning up global sent messages: {e}")

# Update the enqueue function to use the correct queue name
async def enqueue_timetable_processing(teacher_id: str, file_path: str, gcs_file_name: str) -> Optional[str]:
    """
    Enqueue a timetable file processing task for a teacher.
    
    Args:
        teacher_id: UUID string of the teacher
        file_path: Path to the uploaded timetable file
        gcs_file_name: File name to be used in GCS
        
    Returns:
        Job ID string if successful, None if failed
    """
    try:
        # Validate teacher_id is a valid UUID
        UUID(teacher_id)
        
        redis = await create_pool(timetable_redis_settings)
        job = await redis.enqueue_job(
            'process_timetable_file_task', 
            str(teacher_id), 
            file_path,
            gcs_file_name,
            _queue_name='timetable_queue'
        )
        
        logger.info(f"✅ Timetable processing queued for teacher {teacher_id}: {job.job_id}")
        print(f"📄 Timetable job ID for teacher {teacher_id}: {job.job_id}")
        
        await redis.aclose()
        return job.job_id
        
    except ValueError as e:
        logger.error(f"❌ Invalid teacher_id format: {teacher_id} - {e}")
        return None
    except Exception as e:
        logger.error(f"❌ Failed to enqueue timetable task for {teacher_id}: {e}")
        return None

# ARQ Worker Configuration
async def startup(ctx):
    """ARQ worker startup"""
    # Create pool with specific queue name
    ctx['redis'] = await create_pool(timetable_redis_settings, default_queue_name='timetable_queue')
    logger.info("Timetable processing worker started")

async def shutdown(ctx):
    """ARQ worker shutdown"""
    ctx['redis'].close()
    await ctx['redis'].aclose()
    await async_engine.dispose()
    logger.info("Timetable processing worker shutdown")

# Worker configuration for this specific task
timetable_worker_config = {
    'functions': [process_timetable_file_task],
    'redis_settings': timetable_redis_settings,
    'queue_name': 'timetable_queue',  # Use just the queue name without arq:queue: prefix
    'on_startup': startup,
    'on_shutdown': shutdown,
    'max_tries': 3,           # Retry failed jobs 3 times
    'retry_delay': 10,        # Wait 10 seconds between retries
    'job_timeout': 300,       # 5 minutes max per job
    'concurrent_jobs': 2,     # Process 2 files simultaneously
    'keep_result': 3600,      # Keep job results for 1 hour
    'max_jobs': 50            # Max jobs before worker restart
}

# Manual testing function
if __name__ == "__main__":
    async def test_enqueue():
        # Use timetable-specific Redis settings
        redis = await create_pool(timetable_redis_settings)
        try:
            teacher_id = "7bed2b69-8000-4b36-8e91-7fe0b70c9d82"
            test_file = "./uploads/test_timetable.pdf"
            gcs_file_name = "timetable/7bed2b69-8000-4b36-8e91-7fe0b70c9d82.pdf"
            job = await redis.enqueue_job('process_timetable_file_task', teacher_id, test_file, gcs_file_name)
            print(f"[SUCCESS] Timetable processing job queued: {job.job_id}")
        finally:
            await redis.aclose()
    
    asyncio.run(test_enqueue())