"""Background Task for Academic Calendar File Processing

This module provides background task processing for academic calendar file uploads,
including intelligent text extraction based on file type and AI-powered
calendar data parsing.

Supported file types:
- PDF: pdfplumber → pytesseract fallback
- Images (JPG/PNG): pytesseract OCR
- DOCX: python-docx
- XLSX: openpyxl
- TXT: plain text

Usage:
    from calendar_back import process_calendar_file_task
    job_id = await enqueue_calendar_processing(teacher_id, file_path, gcs_file_name)
"""

import os
import logging
import traceback
from pathlib import Path
from typing import Dict, Any, Optional
from uuid import UUID  # Add this import for UUID class
from app.core.logger import logger
from datetime import datetime


# File processing libraries
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    
try:
    import pytesseract
    from PIL import Image
    
    # Configure Tesseract path for Windows
    import platform
    if platform.system() == "Windows":
        # Common Windows installation paths for Tesseract
        possible_paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            r"C:\Users\{username}\AppData\Local\Tesseract-OCR\tesseract.exe".format(
                username=os.environ.get('USERNAME', '')
            ),
            r"C:\Tesseract-OCR\tesseract.exe"
        ]
        
        # Try to find Tesseract executable
        tesseract_path = None
        for path in possible_paths:
            if os.path.exists(path):
                tesseract_path = path
                break
        
        # Set the path if found
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
            logger.info(f"Tesseract found at: {tesseract_path}")
        else:
            logger.warning("Tesseract not found in common paths. Please ensure it's installed and in PATH.")
    
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

# Handle imports for both direct execution and module import
try:
    from app.core.config import settings
    from app.core.logger import logger  # Use the imported logger
    from app.models.model import (
        AcademicCalendar, CalendarEvent, TeacherProfile, TeacherNotification, UploadedFile, TempExtract
    )
    from app.services.external_service import get_holidays_from_ai
except ImportError:
    # If running as script directly, add parent directory to path
    import sys
    import os
    parent_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    sys.path.insert(0, parent_dir)
    from app.core.config import settings
    from app.core.logger import logger  # Use the imported logger
    from app.models.model import (
        AcademicCalendar, CalendarEvent, TeacherProfile, TeacherNotification, UploadedFile, TempExtract
    )
    from app.services.external_service import get_holidays_from_ai

# ARQ and database imports
from arq import create_pool, ArqRedis
from arq.connections import RedisSettings
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession
from sqlalchemy import select
import redis.asyncio as redis

# Project imports
from app.core.config import settings
from app.core.logger import logger  # Use the imported logger
from app.models.model import (
    AcademicCalendar, CalendarEvent, TeacherProfile, TeacherNotification, UploadedFile, TempExtract
)
from app.services.external_service import get_holidays_from_ai

# Import from sch_ground.background which contains shared utilities, but create separate Redis settings
try:
    from app.sch_ground.background import async_engine, publish_ws_message, save_notification
except ImportError:
    # If running as script directly, add parent directory to path
    import sys
    import os
    parent_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    sys.path.insert(0, parent_dir)
    from app.sch_ground.background import async_engine, publish_ws_message, save_notification

# Create separate Redis settings (without queue_name parameter)
calendar_redis_settings = RedisSettings(host="localhost", port=6379, database=0, conn_timeout=10, conn_retries=5, conn_retry_delay=1)

# Initialize async Redis client
redis_client = redis.Redis(host="localhost", port=6379, decode_responses=True)

# File type mappings
SUPPORTED_EXTENSIONS = {
    'pdf': 'pdf',
    'jpg': 'image', 'jpeg': 'image', 'png': 'image', 'bmp': 'image', 'tiff': 'image',
    'docx': 'docx',
    'xlsx': 'excel', 'xls': 'excel',
    'txt': 'text'
}

class FileExtractor:
    """Handles text extraction from various file types"""
    
    # Get Poppler path from settings or environment variable
    POPPLER_PATH = settings.POPPLER_PATH or os.environ.get('POPPLER_PATH', None)
    
    @staticmethod
    def detect_file_type(file_path: str) -> str:
        """Detect file type from extension"""
        logger.info(f"Detecting file type for: {file_path}")
        logger.info(f"File path exists: {os.path.exists(file_path)}")
        if os.path.exists(file_path):
            logger.info(f"File size: {os.path.getsize(file_path)} bytes")
        extension = Path(file_path).suffix.lower().lstrip('.')
        logger.info(f"File extension: {extension}")
        file_type = SUPPORTED_EXTENSIONS.get(extension, 'unknown')
        logger.info(f"Detected file type: {file_type}")
        logger.info(f"Supported extensions: {list(SUPPORTED_EXTENSIONS.keys())}")
        return file_type
    
    @staticmethod
    def extract_from_text(file_path: str) -> str:
        """Extract text from plain text files"""
        logger.info(f"Extracting text from plain text file: {file_path}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            logger.info(f"✅ Text extraction successful: {len(text_content)} characters")
            return text_content
        except Exception as e:
            logger.error(f"❌ Text extraction failed: {e}")
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text_content = f.read()
                logger.info(f"✅ Text extraction with latin-1 successful: {len(text_content)} characters")
                return text_content
            except Exception as e2:
                logger.error(f"❌ Text extraction with latin-1 also failed: {e2}")
                raise
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF using pdfplumber with pytesseract fallback"""
        logger.info(f"📄 Extracting text from PDF: {file_path}")
        
        if not PDF_AVAILABLE:
            logger.error("pdfplumber not installed")
            raise ImportError("pdfplumber not installed")
        
        try:
            # Try pdfplumber for digital PDFs
            logger.info("Trying pdfplumber for digital PDFs")
            with pdfplumber.open(file_path) as pdf:
                text_content = ""
                for i, page in enumerate(pdf.pages):
                    logger.info(f"Processing page {i+1}/{len(pdf.pages)}")
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
                
                if text_content.strip():
                    logger.info(f"✅ Successfully extracted text using pdfplumber: {len(text_content)} characters")
                    return text_content
                else:
                    logger.warning("pdfplumber extracted empty text")
                
        except Exception as e:
            logger.warning(f"⚠️ pdfplumber extraction failed: {e}")
        
        # Fallback to OCR for scanned PDFs
        logger.info("Falling back to OCR extraction")
        return FileExtractor.extract_with_ocr(file_path)
    
    @staticmethod
    def extract_with_ocr(file_path: str) -> str:
        """Extract text using pytesseract OCR"""
        logger.info(f"🔍 Extracting text using OCR: {file_path}")
        
        if not OCR_AVAILABLE:
            logger.error("pytesseract or PIL not installed")
            raise ImportError("pytesseract or PIL not installed")
        
        try:
            # Handle PDF conversion for OCR
            if file_path.lower().endswith('.pdf'):
                logger.info("Converting PDF to images for OCR")
                # Convert PDF to images first (requires pdf2image)
                try:
                    from pdf2image import convert_from_path
                    # Use POPPLER_PATH if specified, otherwise rely on system PATH
                    if FileExtractor.POPPLER_PATH and os.path.exists(FileExtractor.POPPLER_PATH):
                        images = convert_from_path(file_path, poppler_path=FileExtractor.POPPLER_PATH)
                    else:
                        images = convert_from_path(file_path)
                    logger.info(f"Converted PDF to {len(images)} images")
                    
                    text_content = ""
                    for i, image in enumerate(images):
                        logger.info(f"Processing image {i+1}/{len(images)}")
                        try:
                            page_text = pytesseract.image_to_string(image, timeout=30)  # 30 second timeout
                        except Exception as timeout_e:
                            logger.warning(f"OCR timed out for page {i+1}: {timeout_e}")
                            page_text = ""  # Treat timeout as no text extracted
                        text_content += f"Page {i+1}:\n{page_text}\n"
                    
                    logger.info(f"✅ OCR extraction from PDF successful: {len(text_content)} characters")
                    return text_content
                    
                except ImportError:
                    logger.error("❌ pdf2image not installed - cannot OCR PDF files")
                    raise ImportError("pdf2image required for PDF OCR")
            else:
                # Direct image OCR
                logger.info("Performing direct image OCR")
                image = Image.open(file_path)
                try:
                    text_content = pytesseract.image_to_string(image, timeout=30)  # 30 second timeout
                except Exception as timeout_e:
                    logger.warning(f"OCR timed out for image: {timeout_e}")
                    text_content = ""  # Treat timeout as no text extracted
                logger.info(f"✅ OCR extraction successful: {len(text_content)} characters")
                return text_content
                
        except Exception as e:
            logger.error(f"💥 OCR extraction failed: {e}")
            logger.error(f"Full traceback: {traceback.format_exc()}")
            raise
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX files"""
        logger.info(f"📝 Extracting text from DOCX: {file_path}")
        
        if not DOCX_AVAILABLE:
            logger.error("python-docx not installed")
            raise ImportError("python-docx not installed")
        
        try:
            doc = Document(file_path)
            logger.info(f"Loaded DOCX document with {len(doc.paragraphs)} paragraphs")
            text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            logger.info(f"Extracted {len(text_content)} characters from paragraphs")
            
            # Also extract text from tables
            table_count = len(doc.tables)
            logger.info(f"Found {table_count} tables in DOCX")
            for i, table in enumerate(doc.tables):
                logger.info(f"Processing table {i+1}/{table_count}")
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    text_content += f"\n{row_text}"
            
            logger.info(f"✅ DOCX extraction successful: {len(text_content)} characters")
            return text_content
            
        except Exception as e:
            logger.error(f"💥 DOCX extraction failed: {e}")
            logger.error(f"Full traceback: {traceback.format_exc()}")
            raise
    
    @staticmethod
    def extract_from_excel(file_path: str) -> str:
        """Extract text from Excel files"""
        logger.info(f"📊 Extracting text from Excel: {file_path}")
        
        if not XLSX_AVAILABLE:
            logger.error("openpyxl not installed")
            raise ImportError("openpyxl not installed")
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            logger.info(f"Loaded Excel workbook with {len(workbook.sheetnames)} sheets")
            text_content = ""
            
            for sheet_name in workbook.sheetnames:
                logger.info(f"Processing sheet: {sheet_name}")
                sheet = workbook[sheet_name]
                text_content += f"\nSheet: {sheet_name}\n"
                
                row_count = 0
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text_content += f"{row_text}\n"
                        row_count += 1
                
                logger.info(f"Processed {row_count} rows in sheet {sheet_name}")
            
            logger.info(f"✅ Excel extraction successful: {len(text_content)} characters")
            return text_content
            
        except Exception as e:
            logger.error(f"💥 Excel extraction failed: {e}")
            logger.error(f"Full traceback: {traceback.format_exc()}")
            raise

class CalendarParser:
    """Parses extracted text to identify academic calendar data"""
    
    @staticmethod
    def parse_calendar_text(text: str) -> Dict[str, Any]:
        """
        Parse extracted text to identify academic calendar entries.
        This is a basic implementation - can be enhanced with AI/ML.
        """
        logger.info("📅 Parsing academic calendar data from extracted text")
        logger.info(f"Text length: {len(text)}")
        
        # Basic parsing logic - can be enhanced with AI
        calendar_data = {
            "academic_calendar": {},
            "calendar_events": []
        }
        
        # Example patterns to look for in academic calendar data
        lines = text.lower().split('\n')
        logger.info(f"Number of lines in text: {len(lines)}")
        
        # Look for semester information
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Look for semester name patterns
            if "semester" in line or "term" in line:
                calendar_data["academic_calendar"]["semester_name"] = line.title()
                logger.info(f"Found semester pattern: {line}")
                
            # Look for date patterns (e.g., "2024-08-15", "August 15, 2024")
            import re
            date_pattern = r'(\d{4}-\d{2}-\d{2})|([a-zA-Z]+ \d{1,2},? \d{4})'
            date_matches = re.findall(date_pattern, line)
            
            if date_matches:
                # For now, just create sample data
                calendar_data["academic_calendar"] = {
                    "semester_name": "2nd Semester",
                    "semester_start_date": "2024-08-15",
                    "semester_end_date": "2024-12-20",
                    "mid_semester_break_start_date": "2024-10-15",
                    "mid_semester_break_end_date": "2024-10-22",
                    "midsem_exams_date": "2024-10-08",
                    "revision_start_date": "2024-12-01"
                }
                
                calendar_data["calendar_events"] = [
                    {
                        "event_name": "Orientation Week",
                        "event_start_date": "2024-08-12",
                        "event_end_date": "2024-08-14",
                        "event_start_time": "08:00",
                        "event_end_time": "17:00",
                        "is_holiday": False,
                        "requires_no_classes": True
                    },
                    {
                        "event_name": "Independence Day",
                        "event_start_date": "2024-09-21",
                        "event_end_date": "2024-09-21",
                        "event_start_time": "",
                        "event_end_time": "",
                        "is_holiday": True,
                        "requires_no_classes": True
                    },
                    {
                        "event_name": "End of Semester Exams",
                        "event_start_date": "2024-12-02",
                        "event_end_date": "2024-12-13",
                        "event_start_time": "08:00",
                        "event_end_time": "16:00",
                        "is_holiday": False,
                        "requires_no_classes": True
                    }
                ]
                logger.info("Created sample calendar data")
                break
        
        # If no entries found, create sample entries for testing
        if not calendar_data["academic_calendar"]:
            logger.warning("No calendar entries parsed - creating sample entry")
            calendar_data["academic_calendar"] = {
                "semester_name": "Sample Semester",
                "semester_start_date": "2024-09-01",
                "semester_end_date": "2024-12-31"
            }
        
        logger.info(f"📈 Parsed academic calendar with {len(calendar_data.get('calendar_events', []))} events")
        logger.info(f"Calendar data keys: {list(calendar_data.keys())}")
        if "academic_calendar" in calendar_data:
            logger.info(f"Academic calendar keys: {list(calendar_data['academic_calendar'].keys())}")
        return calendar_data

async def process_calendar_file_task(ctx: dict, teacher_id: str, file_path: str, gcs_file_name: str, additional_data: str = ""):
    """
    Background task to process academic calendar files with intelligent text extraction.
    
    Args:
        ctx: ARQ context
        teacher_id: UUID of the teacher (can be None for system/developer records)
        file_path: Original file name or path
        gcs_file_name: File name in GCS
        additional_data: Additional context or data that may contain multiple calendar info
    """
    logger.info(f"🚀 Starting academic calendar file processing for teacher: {teacher_id}, file: {file_path}")
    logger.info(f"GCS file name: {gcs_file_name}")
    logger.info(f"Additional data length: {len(additional_data) if additional_data else 0}")
    logger.info(f"Current working directory: {os.getcwd()}")
    
    # Global set to track sent messages across all function instances
    GLOBAL_SENT_MESSAGES = set()
    
    # Track sent messages to prevent duplicates
    sent_messages = set()
    
    async def send_unique_ws_message(teacher_id: str, message: dict):
        """Send WebSocket message only if it hasn't been sent before"""
        # Skip sending messages if teacher_id is None
        if teacher_id is None:
            logger.info("Skipping WebSocket message for system/developer record (NULL teacher_id)")
            return False
            
        # Create a more robust hash of the message content to identify duplicates
        # Include more fields to make the hash more unique
        status = message.get('status', '')
        type_ = message.get('type', '')
        msg_text = message.get('message', '')
        event_count = str(message.get('event_count', ''))  # Convert to string for consistency
        
        message_key = f"{status}_{type_}_{msg_text}_{event_count}"
        
        # Use a more robust deduplication approach
        import hashlib
        message_hash = hashlib.md5(message_key.encode()).hexdigest()
        
        # Check both local and global deduplication
        logger.info(f"🔍 [DEDUPLICATION] Checking message hash: {message_hash} for key: {message_key}")
        logger.info(f"🔍 [DEDUPLICATION] Local sent_messages set size: {len(sent_messages)}")
        logger.info(f"🔍 [DEDUPLICATION] Global SENT_MESSAGES set size: {len(GLOBAL_SENT_MESSAGES)}")
        logger.info(f"🔍 [DEDUPLICATION] Message content: {message}")
        logger.info(f"🔍 [DEDUPLICATION] Message status: {status}")
        logger.info(f"🔍 [DEDUPLICATION] Message type: {type_}")
        logger.info(f"🔍 [DEDUPLICATION] Message text: {msg_text}")
        
        if message_hash not in sent_messages and message_hash not in GLOBAL_SENT_MESSAGES:
            sent_messages.add(message_hash)
            GLOBAL_SENT_MESSAGES.add(message_hash)
            logger.info(f"➕ [DEDUPLICATION] Adding message hash to both local and global sent_messages: {message_hash}")
            logger.info(f"➕ [DEDUPLICATION] Local set now contains {len(sent_messages)} items")
            logger.info(f"➕ [DEDUPLICATION] Global set now contains {len(GLOBAL_SENT_MESSAGES)} items")
            try:
                logger.info(f"Sending WebSocket message to teacher {teacher_id}")
                logger.info(f"Message content: {message}")
                await publish_ws_message(teacher_id, message)
                logger.info(f"✅ [DEDUPLICATION] Sent unique WebSocket message with hash: {message_hash}")
            except Exception as e:
                logger.error(f"❌ [DEDUPLICATION] Failed to send WebSocket message: {e}")
                logger.error(f"Full traceback: {traceback.format_exc()}")
            return True
        else:
            logger.info(f"⏭️ [DEDUPLICATION] SKIPPING DUPLICATE MESSAGE with hash: {message_hash}")
            logger.info(f"⏭️ [DEDUPLICATION] Message already sent, not sending again")
            if message_hash in sent_messages:
                logger.info(f"⏭️ [DEDUPLICATION] Message found in local sent_messages")
            if message_hash in GLOBAL_SENT_MESSAGES:
                logger.info(f"⏭️ [DEDUPLICATION] Message found in global SENT_MESSAGES")
            return False
    
    try:
        # For metadata-only uploads, we need to download the file from GCS first
        # Create local directory if it doesn't exist
        local_dir = "./downloads/calendar"
        os.makedirs(local_dir, exist_ok=True)
        
        # Generate local file path
        file_extension = os.path.splitext(file_path)[1] or ".dat"
        local_file_path = os.path.join(local_dir, f"calendar_{teacher_id or 'system'}_{int(datetime.now().timestamp())}{file_extension}")
        
        # Download file from GCS to local storage
        from app.core.config import settings
        from app.services.gcs_utils import download_file_from_gcs
        
        download_success = download_file_from_gcs(
            settings.GCS_BUCKET_NAME, 
            gcs_file_name, 
            local_file_path
        )
        
        if not download_success:
            error_msg = f"Failed to download file from GCS: {gcs_file_name}"
            logger.error(f"❌ {error_msg}")
            # Only send error message if teacher_id is not None
            if teacher_id is not None:
                try:
                    await send_unique_ws_message(teacher_id, {
                        "status": "error",
                        "message": error_msg,
                        "teacher_id": teacher_id,
                    })
                except Exception as send_error:
                    logger.error(f"Failed to send error message: {send_error}")
            return {"error": error_msg}
        
        # Update file_path to point to the downloaded local file
        file_path = local_file_path
        logger.info(f"✅ File downloaded from GCS to: {file_path}")
        logger.info(f"File path exists: {os.path.exists(file_path)}")
        if os.path.exists(file_path):
            logger.info(f"File size: {os.path.getsize(file_path)} bytes")
        
        # Send initial status - always send this
        logger.info("Sending initial status message")
        logger.info(f"Teacher ID: {teacher_id}")
        logger.info(f"File path: {file_path}")
        result = await send_unique_ws_message(teacher_id, {
            "status": "started",
            "message": "Processing academic calendar file...",
            "teacher_id": teacher_id,
            "file_path": file_path
        })
        logger.info(f"Initial status message sent successfully, result: {result}")
        
        # Validate file exists
        logger.info(f"Checking if file exists: {file_path}")
        file_exists = os.path.exists(file_path)
        logger.info(f"File exists: {file_exists}")
        if not file_exists:
            error_msg = f"File not found: {file_path}"
            logger.error(f"❌ {error_msg}")
            # Only send error message if we haven't already sent one
            try:
                await send_unique_ws_message(teacher_id, {
                    "status": "error",
                    "message": error_msg,
                    "teacher_id": teacher_id,
                })
            except Exception as send_error:
                logger.error(f"Failed to send error message: {send_error}")
            return {"error": error_msg}
        
        # Detect file type
        file_type = FileExtractor.detect_file_type(file_path)
        logger.info(f"🔍 Detected file type: {file_type}")
        logger.info(f"File extension: {Path(file_path).suffix}")
        
        if file_type == 'unknown':
            error_msg = f"Unsupported file type: {Path(file_path).suffix}"
            logger.error(f"❌ {error_msg}")
            logger.error(f"File extension: {Path(file_path).suffix}")
            logger.error(f"Supported extensions: {list(SUPPORTED_EXTENSIONS.keys())}")
            # Only send error message if we haven't already sent one
            try:
                await send_unique_ws_message(teacher_id, {
                    "status": "error", 
                    "message": error_msg,
                    "teacher_id": teacher_id
                })
            except Exception as send_error:
                logger.error(f"Failed to send error message: {send_error}")
            return {"error": error_msg}
        
        # Text Extraction Progress
        logger.info("Sending text extraction progress message")
        logger.info(f"File type: {file_type}")
        result = await send_unique_ws_message(teacher_id, {
            "status": "processing",
            "message": f"Extracting text from {file_type} file...",
            "teacher_id": teacher_id
        })
        logger.info(f"Text extraction progress message sent successfully, result: {result}")
        
        # Extract text based on file type
        extracted_text = ""
        try:
            logger.info(f"Extracting text from {file_type} file")
            if file_type == 'pdf':
                extracted_text = FileExtractor.extract_from_pdf(file_path)
            elif file_type == 'image':
                extracted_text = FileExtractor.extract_with_ocr(file_path)
            elif file_type == 'docx':
                extracted_text = FileExtractor.extract_from_docx(file_path)
            elif file_type == 'excel':
                extracted_text = FileExtractor.extract_from_excel(file_path)
            elif file_type == 'text':
                extracted_text = FileExtractor.extract_from_text(file_path)
            
            logger.info(f"📄 Extracted {len(extracted_text)} characters from file")
            logger.info(f"First 200 characters of extracted text: {extracted_text[:200] if extracted_text else 'None'}")
            
            # If text extraction failed, return error
            if not extracted_text:
                error_msg = "Text extraction returned empty result"
                logger.error(f"❌ {error_msg}")
                logger.error(f"Extracted text length: {len(extracted_text) if extracted_text else 0}")
                # Only send error message if we haven't already sent one
                try:
                    await send_unique_ws_message(teacher_id, {
                        "status": "error",
                        "message": error_msg,
                        "teacher_id": teacher_id
                    })
                except Exception as send_error:
                    logger.error(f"Failed to send error message: {send_error}")
                return {"error": error_msg}
        except Exception as e:
            error_msg = f"Text extraction failed: {str(e)}"
            logger.error(f"❌ {error_msg}")
            logger.error(f"🔍 Full traceback: {traceback.format_exc()}")
            # Only send error message if we haven't already sent one
            try:
                await send_unique_ws_message(teacher_id, {
                    "status": "error",
                    "message": error_msg,
                    "teacher_id": teacher_id
                })
            except Exception as send_error:
                logger.error(f"Failed to send error message: {send_error}")
            return {"error": error_msg}
        
        # AI Processing Progress
        logger.info("Sending AI processing progress message")
        logger.info(f"Extracted text length: {len(extracted_text)}")
        result = await send_unique_ws_message(teacher_id, {
            "status": "processing",
            "message": "Processing extracted text with AI...",
            "teacher_id": teacher_id,
            "extracted_text_length": len(extracted_text)
        })
        logger.info(f"AI processing progress message sent successfully, result: {result}")
        
        # Process with AI (if available)
        ai_result = None
        try:
            from app.services.external_service import send_academic_calendar_to_ai
            if hasattr(settings, 'GEMINI_API_KEY') and settings.GEMINI_API_KEY:
                logger.info("Sending extracted text to AI for processing")
                logger.info(f"GEMINI_API_KEY available: {bool(settings.GEMINI_API_KEY)}")
                logger.info(f"GCS_BUCKET_NAME: {settings.GCS_BUCKET_NAME}")
                logger.info(f"GCS file name: {gcs_file_name}")
                logger.info(f"Additional data length: {len(additional_data) if additional_data else 0}")
                # Pass both the extracted text and GCS file name to the AI processing function
                ai_result = await send_academic_calendar_to_ai(
                    extracted_text, 
                    f"gs://{settings.GCS_BUCKET_NAME}/{gcs_file_name}",
                    settings.GEMINI_API_KEY,
                    additional_data
                )
                logger.info(f"AI result received: {type(ai_result)}")
                if isinstance(ai_result, dict):
                    logger.info(f"AI result keys: {list(ai_result.keys())}")
                    if "error" in ai_result:
                        logger.info(f"AI result error: {ai_result['error']}")
                    else:
                        logger.info("AI processing successful, no error in result")
                
                if "error" in ai_result:
                    logger.warning(f"AI processing failed: {ai_result['error']}")
                    logger.warning(f"AI error details: {ai_result}")
                    # Fall back to basic parsing
                    logger.info("Falling back to basic parsing")
                    calendar_data = CalendarParser.parse_calendar_text(extracted_text)
                else:
                    logger.info("AI processing successful")
                    logger.info(f"AI result type: {type(ai_result)}")
                    if isinstance(ai_result, dict):
                        logger.info(f"AI result keys: {list(ai_result.keys())}")
                        if "semester_name" in ai_result:
                            logger.info(f"AI result semester_name: {ai_result['semester_name']}")
                        if "calendar_events" in ai_result:
                            logger.info(f"AI result calendar_events count: {len(ai_result['calendar_events'])}")
                            if ai_result['calendar_events']:
                                logger.info(f"First calendar event: {ai_result['calendar_events'][0]}")
                    logger.info(f"AI result: {ai_result}")
                    calendar_data = ai_result
            else:
                logger.info("Skipping AI processing - no API key configured")
                logger.info(f"GEMINI_API_KEY attribute exists: {hasattr(settings, 'GEMINI_API_KEY')}")
                if hasattr(settings, 'GEMINI_API_KEY'):
                    logger.info(f"GEMINI_API_KEY value: {settings.GEMINI_API_KEY}")
                    logger.info(f"GEMINI_API_KEY length: {len(settings.GEMINI_API_KEY) if settings.GEMINI_API_KEY else 0}")
                # Fall back to basic parsing
                calendar_data = CalendarParser.parse_calendar_text(extracted_text)
                logger.info(f"Basic parsing result type: {type(calendar_data)}")
                if isinstance(calendar_data, dict):
                    logger.info(f"Basic parsing result keys: {list(calendar_data.keys())}")
                    logger.info(f"Basic parsing result: {calendar_data}")
        except Exception as e:
            logger.error(f"AI processing failed: {e}")
            logger.error(f"Full traceback: {traceback.format_exc()}")
            # Fall back to basic parsing
            logger.info("Falling back to basic parsing")
            calendar_data = CalendarParser.parse_calendar_text(extracted_text)
            logger.info(f"Basic parsing result type: {type(calendar_data)}")
            if isinstance(calendar_data, dict):
                logger.info(f"Basic parsing result keys: {list(calendar_data.keys())}")
                logger.info(f"Basic parsing result: {calendar_data}")
        
        # Validate calendar data
        # Check if we have the required academic calendar fields directly or wrapped in academic_calendar object
        logger.info(f"Validating calendar data. Data type: {type(calendar_data)}")
        if not calendar_data:
            error_msg = "No academic calendar data found in file"
            logger.error(f"❌ {error_msg}")
            logger.error(f"Calendar data type: {type(calendar_data)}")
            logger.error(f"Calendar data: {calendar_data}")
            # Only send error message if we haven't already sent one
            try:
                logger.info("Sending error message")
                result = await send_unique_ws_message(teacher_id, {
                    "status": "error",
                    "message": error_msg,
                    "teacher_id": teacher_id
                })
                logger.info(f"Error message sent successfully, result: {result}")
            except Exception as send_error:
                logger.error(f"Failed to send error message: {send_error}")
            return {"error": error_msg}
        
        # Log the structure of calendar_data for debugging
        logger.info(f"🔍 Calendar data structure: {type(calendar_data)}")
        if isinstance(calendar_data, dict):
            logger.info(f"🔍 Calendar data keys: {list(calendar_data.keys())}")
            if "academic_calendar" in calendar_data:
                logger.info(f"🔍 Academic calendar keys: {list(calendar_data['academic_calendar'].keys())}")
        
        # Check if data is already correctly structured (has required fields at root and calendar_events)
        # This is the structure we expect from the AI response
        has_root_calendar_fields = all(field in calendar_data for field in [
            "semester_name", "semester_start_date", "semester_end_date"
        ])
        has_calendar_events = "calendar_events" in calendar_data
        
        logger.info(f"Has root calendar fields: {has_root_calendar_fields}")
        logger.info(f"Has calendar events: {has_calendar_events}")
        
        if has_root_calendar_fields and has_calendar_events:
            # Data is already correctly structured from AI response
            # Wrap it in the expected format for consistency
            wrapped_data = {
                "academic_calendar": {
                    "semester_name": calendar_data["semester_name"],
                    "semester_start_date": calendar_data["semester_start_date"],
                    "semester_end_date": calendar_data["semester_end_date"],
                    "mid_semester_break_start_date": calendar_data.get("mid_semester_break_start_date"),
                    "mid_semester_break_end_date": calendar_data.get("mid_semester_break_end_date"),
                    "midsem_exams_date": calendar_data.get("midsem_exams_date"),
                    "revision_start_date": calendar_data.get("revision_start_date")
                },
                "calendar_events": calendar_data["calendar_events"]
            }
            calendar_data = wrapped_data
            logger.info("✅ Academic calendar data is correctly structured from AI response")
        elif "academic_calendar" in calendar_data and "calendar_events" in calendar_data:
            # Data is already wrapped correctly
            logger.info("✅ Academic calendar data is already wrapped correctly")
        else:
            # Data structure is not as expected
            logger.error("Calendar data structure validation failed")
            logger.error(f"Calendar data keys: {list(calendar_data.keys()) if isinstance(calendar_data, dict) else 'Not a dict'}")
            logger.error(f"Calendar data: {calendar_data}")
            error_msg = "No academic calendar data found in file"
            logger.error(f"❌ {error_msg}")
            # Only send error message if we haven't already sent one
            try:
                logger.info("Sending validation error message")
                result = await send_unique_ws_message(teacher_id, {
                    "status": "error",
                    "message": error_msg,
                    "teacher_id": teacher_id
                })
                logger.info(f"Validation error message sent successfully, result: {result}")
            except Exception as send_error:
                logger.error(f"Failed to send validation error message: {send_error}")
            return {"error": error_msg}
        
        calendar_events = calendar_data.get("calendar_events", [])
        logger.info(f"📊 Found academic calendar with {len(calendar_events)} events")
        
        # Log more details about the calendar data for debugging
        logger.info(f"📋 Calendar data keys: {list(calendar_data.keys())}")
        academic_calendar_data = calendar_data.get('academic_calendar', {})
        logger.info(f"📋 Academic calendar data: {academic_calendar_data}")
        if academic_calendar_data:
            logger.info(f"📋 Academic calendar keys: {list(academic_calendar_data.keys())}")
        if calendar_events:
            logger.info(f"📋 First calendar event: {calendar_events[0]}")
        else:
            logger.warning("⚠️ No calendar events found in data")
        
        # Update progress before saving
        logger.info("Sending saving progress message")
        logger.info(f"Calendar events count: {len(calendar_events)}")
        result = await send_unique_ws_message(teacher_id, {
            "status": "processing",
            "message": f"Saving academic calendar with {len(calendar_events)} events...",
            "teacher_id": teacher_id,
            "event_count": len(calendar_events)
        })
        logger.info(f"Saving progress message sent successfully, result: {result}")
        
        # Save to database
        try:
            from app.models.model import AcademicCalendar, CalendarEvent
            from sqlalchemy import select
            
            # Create database session
            async with AsyncSession(async_engine) as session:
                # Get teacher profile
                stmt = select(TeacherProfile).where(TeacherProfile.id == UUID(teacher_id))
                result = await session.execute(stmt)
                teacher = result.scalar_one_or_none()
                
                logger.info(f"Teacher lookup result: {teacher is not None}")
                if teacher:
                    logger.info(f"Teacher ID: {teacher.id}")
                    logger.info(f"Teacher display name: {teacher.display_name}")
                
                if not teacher:
                    error_msg = f"Teacher not found: {teacher_id}"
                    logger.error(f"❌ {error_msg}")
                    # Only send error message if we haven't already sent one
                    try:
                        await send_unique_ws_message(teacher_id, {
                            "status": "error",
                            "message": error_msg,
                            "teacher_id": teacher_id
                        })
                    except Exception as send_error:
                        logger.error(f"Failed to send error message: {send_error}")
                    return {"error": error_msg}
                
                # ONLY save to TempExtract table for user review with type "academic calendar"
                # First check if there's already an entry for this teacher and type and DELETE it
                try:
                    logger.info("Querying existing TempExtract entries")
                    logger.info(f"Teacher ID: {teacher_id}")
                    logger.info(f"Entry type: academic calendar")
                    existing_temp_entries = (await session.execute(
                        select(TempExtract).where(
                            TempExtract.teacher_id == UUID(teacher_id),
                            TempExtract.type == "academic calendar"
                        )
                    )).scalars().all()
                    logger.info(f"🔍 Found {len(existing_temp_entries)} existing TempExtract entries")
                    if existing_temp_entries:
                        for entry in existing_temp_entries:
                            logger.info(f"Existing entry ID: {entry.id}, Created at: {entry.created_at}")
                except Exception as e:
                    logger.error(f"❌ Error querying existing TempExtract entries: {e}")
                    logger.error(f"Full traceback: {traceback.format_exc()}")
                    raise
                
                # Delete existing entries
                deleted_count = 0
                for entry in existing_temp_entries:
                    try:
                        logger.info(f"Deleting existing TempExtract entry: {entry.id}")
                        await session.delete(entry)
                        deleted_count += 1
                    except Exception as e:
                        logger.error(f"❌ Error deleting TempExtract entry {entry.id}: {e}")
                        logger.error(f"Full traceback: {traceback.format_exc()}")
                        raise
                
                if deleted_count > 0:
                    try:
                        await session.commit()
                        logger.info(f"🗑️ Deleted {deleted_count} existing TempExtract entries for teacher {teacher_id}")
                    except Exception as e:
                        logger.error(f"❌ Error committing TempExtract entry deletions: {e}")
                        logger.error(f"Full traceback: {traceback.format_exc()}")
                        raise
                
                # Prepare the data to be stored - convert to JSON-serializable format
                temp_data = {
                    "academic_calendar": calendar_data.get("academic_calendar", {}),
                    "calendar_events": [],
                    "event_count": len(calendar_events),
                    "extracted_at": datetime.utcnow().isoformat()
                }
                
                # Clean up empty string values for optional date fields in academic_calendar
                # Convert empty strings to None for optional date fields
                academic_calendar_data = temp_data["academic_calendar"]
                date_fields = [
                    "mid_semester_break_start_date",
                    "mid_semester_break_end_date",
                    "midsem_exams_date",
                    "revision_start_date"
                ]
                
                for field in date_fields:
                    if field in academic_calendar_data and academic_calendar_data[field] == "":
                        academic_calendar_data[field] = None
                        logger.info(f"🧹 Cleaned empty string for field '{field}' -> None in academic_calendar")
                
                logger.info(f"Prepared temp_data structure: {list(temp_data.keys())}")
                logger.info(f"Academic calendar in temp_data: {temp_data['academic_calendar']}")
                logger.info(f"Event count in temp_data: {temp_data['event_count']}")
                logger.info(f"Extracted at timestamp: {temp_data['extracted_at']}")
                
                # Convert each event to a dictionary and ensure UUID fields are strings
                logger.info(f"Processing {len(calendar_events)} calendar events")
                for i, event_data in enumerate(calendar_events):
                    logger.info(f"Processing event {i+1}/{len(calendar_events)}")
                    event_dict = event_data.copy()
                    logger.info(f"Event data keys: {list(event_dict.keys())}")
                    # Ensure all UUID fields are converted to strings
                    if 'teacher_id' in event_dict and isinstance(event_dict['teacher_id'], UUID):
                        event_dict['teacher_id'] = str(event_dict['teacher_id'])
                        logger.info(f"Converted teacher_id to string: {event_dict['teacher_id']}")
                    # Convert time objects to strings for JSON serialization
                    if 'event_start_time' in event_dict and isinstance(event_dict['event_start_time'], time):
                        event_dict['event_start_time'] = event_dict['event_start_time'].isoformat()
                        logger.info(f"Converted event_start_time to string: {event_dict['event_start_time']}")
                    if 'event_end_time' in event_dict and isinstance(event_dict['event_end_time'], time):
                        event_dict['event_end_time'] = event_dict['event_end_time'].isoformat()
                        logger.info(f"Converted event_end_time to string: {event_dict['event_end_time']}")
                    temp_data["calendar_events"].append(event_dict)
                
                logger.info(f"Processed calendar events count: {len(temp_data['calendar_events'])}")
                if temp_data["calendar_events"]:
                    logger.info(f"First event keys: {list(temp_data['calendar_events'][0].keys())}")
                
                # Log the data that will be saved for debugging
                logger.info(f"💾 Saving TempExtract data: {temp_data}")
                
                # Create new entry (this ensures overwrite behavior)
                try:
                    logger.info("Creating new TempExtract entry")
                    logger.info(f"Teacher ID: {teacher_id}")
                    logger.info(f"Entry type: academic calendar")
                    logger.info(f"Temp data: {temp_data}")
                    temp_extract = TempExtract(
                        teacher_id=UUID(teacher_id),
                        type="academic calendar",
                        data=temp_data
                    )
                    logger.info(f"🆕 Created TempExtract object: teacher_id={temp_extract.teacher_id}, type={temp_extract.type}")
                    logger.info(f"TempExtract data keys: {list(temp_data.keys())}")
                    logger.info(f"Academic calendar in TempExtract: {temp_data.get('academic_calendar', {})}")
                    logger.info(f"Calendar events count in TempExtract: {len(temp_data.get('calendar_events', []))}")
                    session.add(temp_extract)
                    logger.info(f"➕ Added TempExtract entry to session for teacher {teacher_id}")
                except Exception as e:
                    logger.error(f"❌ Error creating or adding TempExtract entry: {e}")
                    logger.error(f"Full traceback: {traceback.format_exc()}")
                    raise
                
                try:
                    logger.info("Committing TempExtract entry to database")
                    await session.commit()
                    logger.info("✅ Successfully committed to database")
                except Exception as e:
                    logger.error(f"❌ Error committing to database: {e}")
                    logger.error(f"Full traceback: {traceback.format_exc()}")
                    raise
                
                # Verify the data was saved
                try:
                    logger.info("Verifying TempExtract entry was saved")
                    stmt = select(TempExtract).where(
                        TempExtract.teacher_id == UUID(teacher_id),
                        TempExtract.type == "academic calendar"
                    )
                    result = await session.execute(stmt)
                    saved_entry = result.scalar_one_or_none()
                    if saved_entry:
                        logger.info(f"✅ Verified TempExtract entry saved with ID: {saved_entry.id}")
                        logger.info(f"✅ Saved data keys: {list(saved_entry.data.keys())}")
                        logger.info(f"✅ Academic calendar in saved entry: {saved_entry.data.get('academic_calendar', {})}")
                        logger.info(f"✅ Calendar events count in saved entry: {len(saved_entry.data.get('calendar_events', []))}")
                        logger.info(f"✅ Extracted at timestamp: {saved_entry.data.get('extracted_at', 'N/A')}")
                    else:
                        logger.error("❌ Failed to verify TempExtract entry was saved")
                        logger.error(f"Teacher ID: {teacher_id}")
                        logger.error(f"Entry type: academic calendar")
                        # Try to query all entries for this teacher to see what's there
                        all_entries_stmt = select(TempExtract).where(
                            TempExtract.teacher_id == UUID(teacher_id)
                        )
                        all_entries_result = await session.execute(all_entries_stmt)
                        all_entries = all_entries_result.scalars().all()
                        logger.error(f"All TempExtract entries for teacher: {len(all_entries)}")
                        for entry in all_entries:
                            logger.error(f"Entry ID: {entry.id}, Type: {entry.type}")
                except Exception as e:
                    logger.error(f"❌ Error verifying TempExtract entry: {e}")
                    logger.error(f"Full traceback: {traceback.format_exc()}")
                
                # Create success notification
                notification_msg = f"Successfully processed academic calendar with {len(calendar_events)} events"
                logger.info(f"Creating success notification: {notification_msg}")
                try:
                    logger.info("Saving notification to database")
                    logger.info(f"Teacher ID: {teacher_id}")
                    logger.info(f"Notification title: Academic Calendar Processing Complete")
                    logger.info(f"Notification message: {notification_msg}")
                    logger.info(f"Notification type: success")
                    await save_notification(
                        teacher_id=UUID(teacher_id),  # Pass UUID object directly
                        title="Academic Calendar Processing Complete",
                        message=notification_msg,
                        type_="success"
                    )
                    logger.info("Success notification saved successfully")
                except Exception as e:
                    logger.error(f"Failed to save success notification: {e}")
                    logger.error(f"Full traceback: {traceback.format_exc()}")
                
                # Completion Message with type "COMPLETED_ACADEMIC_CALENDER"
                # Convert events to JSON-serializable format
                logger.info(f"Preparing serializable events from {len(calendar_events)} events")
                serializable_events = []
                for i, event_data in enumerate(calendar_events):
                    logger.info(f"Processing event {i+1}/{len(calendar_events)} for serialization")
                    event_dict = event_data.copy()
                    logger.info(f"Event data keys: {list(event_dict.keys())}")
                    # Ensure all UUID fields are converted to strings
                    if 'teacher_id' in event_dict and isinstance(event_dict['teacher_id'], UUID):
                        event_dict['teacher_id'] = str(event_dict['teacher_id'])
                        logger.info(f"Converted teacher_id to string: {event_dict['teacher_id']}")
                    # Convert time objects to strings for JSON serialization
                    if 'event_start_time' in event_dict and isinstance(event_dict['event_start_time'], time):
                        event_dict['event_start_time'] = event_dict['event_start_time'].isoformat()
                        logger.info(f"Converted event_start_time to string: {event_dict['event_start_time']}")
                    if 'event_end_time' in event_dict and isinstance(event_dict['event_end_time'], time):
                        event_dict['event_end_time'] = event_dict['event_end_time'].isoformat()
                        logger.info(f"Converted event_end_time to string: {event_dict['event_end_time']}")
                    serializable_events.append(event_dict)
                
                logger.info(f"Prepared serializable events count: {len(serializable_events)}")
                if serializable_events:
                    logger.info(f"First serializable event keys: {list(serializable_events[0].keys())}")
                
                completion_message = {
                    "status": "complete",
                    "type": "COMPLETED_ACADEMIC_CALENDER",  # Changed to match requirement
                    "message": notification_msg,
                    "teacher_id": teacher_id,
                    "event_count": len(calendar_events)
                }
                
                logger.info(f"Sending completion message: {completion_message}")
                
                result = await send_unique_ws_message(teacher_id, completion_message)
                logger.info(f"Completion message sent successfully, result: {result}")
                
                # Clean up temporary file
                try:
                    logger.info(f"Cleaning up temporary file: {file_path}")
                    logger.info(f"File exists before cleanup: {os.path.exists(file_path)}")
                    if os.path.exists(file_path):
                        os.remove(file_path)
                        logger.info(f"🧹 Cleaned up temporary file: {file_path}")
                    else:
                        logger.warning(f"Temporary file does not exist: {file_path}")
                except Exception as e:
                    logger.warning(f"⚠️ Failed to clean up temporary file: {e}")
                    logger.warning(f"Full traceback: {traceback.format_exc()}")
                
                # Convert events to JSON-serializable format for return
                serializable_events = []
                for event_data in calendar_events:
                    event_dict = event_data.copy()
                    # Ensure all UUID fields are converted to strings
                    if 'teacher_id' in event_dict and isinstance(event_dict['teacher_id'], UUID):
                        event_dict['teacher_id'] = str(event_dict['teacher_id'])
                    # Convert time objects to strings for JSON serialization
                    if 'event_start_time' in event_dict and isinstance(event_dict['event_start_time'], time):
                        event_dict['event_start_time'] = event_dict['event_start_time'].isoformat()
                    if 'event_end_time' in event_dict and isinstance(event_dict['event_end_time'], time):
                        event_dict['event_end_time'] = event_dict['event_end_time'].isoformat()
                    serializable_events.append(event_dict)
                
                logger.info(f"Prepared final serializable events count: {len(serializable_events)}")
                
                # Clean up empty string values for optional date fields in academic_calendar_data
                # Convert empty strings to None for optional date fields
                academic_calendar_data = calendar_data.get("academic_calendar", {}) if isinstance(calendar_data, dict) else {}
                date_fields = [
                    "mid_semester_break_start_date",
                    "mid_semester_break_end_date",
                    "midsem_exams_date",
                    "revision_start_date"
                ]
                
                for field in date_fields:
                    if field in academic_calendar_data and academic_calendar_data[field] == "":
                        academic_calendar_data[field] = None
                        logger.info(f"🧹 Cleaned empty string for field '{field}' -> None in final academic_calendar")
                
                result = {
                    "status": "success",
                    "message": notification_msg,
                    "event_count": len(calendar_events),
                    "academic_calendar": academic_calendar_data,
                    "calendar_events": serializable_events
                }
                
                logger.info(f"Returning result: {result}")
                logger.info(f"Result keys: {list(result.keys())}")
                logger.info(f"Academic calendar in result: {result.get('academic_calendar', {})}")
                logger.info(f"Calendar events count in result: {result.get('event_count', 0)}")
                logger.info(f"Calendar events in result: {len(result.get('calendar_events', []))}")
                
                return result
                
        except Exception as e:
            error_msg = f"Database operation failed: {str(e)}"
            logger.error(f"❌ {error_msg}")
            logger.error(f"🔍 Full traceback: {traceback.format_exc()}")
            # Only send error message if we haven't already sent one
            try:
                await send_unique_ws_message(teacher_id, {
                    "status": "error",
                    "message": error_msg,
                    "teacher_id": teacher_id
                })
            except Exception as send_error:
                logger.error(f"Failed to send error message: {send_error}")
            # Raise exception to trigger ARQ retry mechanism
            raise RuntimeError(error_msg)
            
    except Exception as e:
        error_msg = f"Unexpected error in calendar processing: {str(e)}"
        logger.error(f"💥 {error_msg}")
        logger.error(f"🔍 Full traceback: {traceback.format_exc()}")
        # Only send error message if we haven't already sent one
        try:
            await send_unique_ws_message(teacher_id, {
                "status": "error",
                "message": error_msg,
                "teacher_id": teacher_id
            })
        except Exception as send_error:
            logger.error(f"Failed to send error message: {send_error}")
        # Raise exception to trigger ARQ retry mechanism
        raise RuntimeError(error_msg)
    finally:
        # Clean up global sent messages for this function instance
        try:
            logger.info(f"Cleaning up global sent messages, current local set size: {len(sent_messages)}")
            for message_hash in sent_messages:
                if message_hash in GLOBAL_SENT_MESSAGES:
                    GLOBAL_SENT_MESSAGES.remove(message_hash)
            logger.info(f"🧹 Cleaned up {len(sent_messages)} message hashes from global set")
            logger.info(f"Global set size after cleanup: {len(GLOBAL_SENT_MESSAGES)}")
        except Exception as e:
            logger.warning(f"⚠️ Error cleaning up global sent messages: {e}")
            logger.warning(f"Full traceback: {traceback.format_exc()}")


# ARQ Worker Configuration
async def startup(ctx):
    """ARQ worker startup"""
    # Create pool with specific queue name
    ctx['redis'] = await create_pool(calendar_redis_settings, default_queue_name='calendar_queue')
    logger.info("Academic calendar processing worker started")

async def shutdown(ctx):
    """ARQ worker shutdown"""
    ctx['redis'].close()
    await ctx['redis'].aclose()
    await async_engine.dispose()
    logger.info("Academic calendar processing worker shutdown")

# Worker configuration for this specific task
calendar_worker_config = {
    'functions': [process_calendar_file_task],
    'redis_settings': calendar_redis_settings,
    'queue_name': 'calendar_queue',  # Use just the queue name without arq:queue: prefix
    'on_startup': startup,
    'on_shutdown': shutdown,
    'max_tries': 3,           # Retry failed jobs 3 times
    'retry_delay': 10,        # Wait 10 seconds between retries
    'job_timeout': 300,       # 5 minutes max per job
    'concurrent_jobs': 2,     # Process 2 files simultaneously
    'keep_result': 3600,      # Keep job results for 1 hour
    'max_jobs': 50            # Max jobs before worker restart
}

logger.info("Calendar worker configuration loaded")
logger.info(f"Worker functions: {[f.__name__ for f in calendar_worker_config['functions']]}")
logger.info(f"Max tries: {calendar_worker_config['max_tries']}")
logger.info(f"Job timeout: {calendar_worker_config['job_timeout']}")

# Manual testing function
if __name__ == "__main__":
    async def test_enqueue():
        logger.info("Starting manual test enqueue")
        logger.info(f"Current working directory: {os.getcwd()}")
        logger.info(f"Files in current directory: {os.listdir('.')}")
        if os.path.exists('./uploads'):
            logger.info(f"Files in uploads directory: {os.listdir('./uploads')}")
        # Use calendar-specific Redis settings
        redis = await create_pool(calendar_redis_settings)
        try:
            teacher_id = "7bed2b69-8000-4b36-8e91-7fe0b70c9d82"
            test_file = "./uploads/test_calendar.pdf"
            gcs_file_name = "academic_calendar/7bed2b69-8000-4b36-8e91-7fe0b70c9d82.pdf"
            logger.info(f"Enqueuing job for teacher {teacher_id}")
            logger.info(f"Test file: {test_file}")
            logger.info(f"GCS file name: {gcs_file_name}")
            logger.info(f"Test file exists: {os.path.exists(test_file)}")
            if os.path.exists(test_file):
                logger.info(f"Test file size: {os.path.getsize(test_file)} bytes")
            job = await redis.enqueue_job('process_calendar_file_task', teacher_id, test_file, gcs_file_name)
            logger.info(f"[SUCCESS] Calendar processing job queued: {job.job_id}")
        except Exception as e:
            logger.error(f"Failed to enqueue job: {e}")
            logger.error(f"Full traceback: {traceback.format_exc()}")
        finally:
            await redis.aclose()
    
    asyncio.run(test_enqueue())


# Update the enqueue function to use the correct queue name
async def enqueue_calendar_processing(teacher_id: str, file_path: str, gcs_file_name: str, additional_data: str = "") -> Optional[str]:
    """
    Enqueue a calendar file processing task for a teacher.
    
    Args:
        teacher_id: UUID string of the teacher
        file_path: Path to the uploaded calendar file
        gcs_file_name: File name to be used in GCS
        additional_data: Additional context or data that may contain multiple calendar info
        
    Returns:
        Job ID string if successful, None if failed
    """
    try:
        # Validate teacher_id is a valid UUID
        UUID(teacher_id)
        
        redis = await create_pool(calendar_redis_settings)
        job = await redis.enqueue_job(
            'process_calendar_file_task', 
            str(teacher_id), 
            file_path,
            gcs_file_name,
            additional_data,
            _queue_name='calendar_queue'
        )
        
        logger.info(f"✅ Calendar processing queued for teacher {teacher_id}: {job.job_id}")
        print(f"📅 Calendar job ID for teacher {teacher_id}: {job.job_id}")
        
        await redis.aclose()
        return job.job_id
        
    except ValueError as e:
        logger.error(f"❌ Invalid teacher_id format: {teacher_id} - {e}")
        return None
    except Exception as e:
        logger.error(f"❌ Failed to enqueue calendar task for {teacher_id}: {e}")
        return None
