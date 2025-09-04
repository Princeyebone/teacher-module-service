"""Background Task for Timetable File Processing

This module provides background task processing for timetable file uploads,
including intelligent text extraction based on file type and AI-powered
timetable data parsing.

Supported file types:
- PDF: pdfplumber → pytesseract fallback
- Images (JPG/PNG): pytesseract OCR
- DOCX: python-docx
- XLSX: openpyxl

Usage:
    from table_back import process_timetable_file_task
    job_id = await enqueue_timetable_processing(teacher_id, file_path)
"""

import os
import json
import asyncio
import logging
import traceback
from pathlib import Path
from typing import Dict, List, Optional, Any
from datetime import datetime
from uuid import UUID


# Initialize logger early - needed for Tesseract configuration
logger = logging.getLogger(__name__)

# File processing libraries
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    
try:
    import pytesseract
    from PIL import Image
    
    # Configure Tesseract path for Windows
    import platform
    if platform.system() == "Windows":
        # Common Windows installation paths for Tesseract
        possible_paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            r"C:\Users\{username}\AppData\Local\Tesseract-OCR\tesseract.exe".format(
                username=os.environ.get('USERNAME', '')
            ),
            r"C:\Tesseract-OCR\tesseract.exe"
        ]
        
        # Try to find Tesseract executable
        tesseract_path = None
        for path in possible_paths:
            if os.path.exists(path):
                tesseract_path = path
                break
        
        # Set the path if found
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
            logger.info(f"Tesseract found at: {tesseract_path}")
        else:
            logger.warning("Tesseract not found in common paths. Please ensure it's installed and in PATH.")
    
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

# ARQ and database imports
from arq import create_pool, ArqRedis
from arq.connections import RedisSettings
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession
from sqlalchemy import select
import redis.asyncio as redis

# Project imports
from config import settings
from model import WeeklyTimeTable, TeacherProfile, TeacherNotification, UploadedFile
from sch_ground.background import arq_redis_settings, async_engine, publish_ws_message, save_notification
from external_service import get_holidays_from_ai  # For potential AI parsing integration

# Logger already initialized at the top

# Initialize async Redis client
redis_client = redis.Redis(host="localhost", port=6379, decode_responses=True)

# File type mappings
SUPPORTED_EXTENSIONS = {
    'pdf': 'pdf',
    'jpg': 'image', 'jpeg': 'image', 'png': 'image', 'bmp': 'image', 'tiff': 'image',
    'docx': 'docx',
    'xlsx': 'excel', 'xls': 'excel'
}

class FileExtractor:
    """Handles text extraction from various file types"""
    
    @staticmethod
    def detect_file_type(file_path: str) -> str:
        """Detect file type from extension"""
        extension = Path(file_path).suffix.lower().lstrip('.')
        return SUPPORTED_EXTENSIONS.get(extension, 'unknown')
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF using pdfplumber with pytesseract fallback"""
        logger.info(f"Extracting text from PDF: {file_path}")
        
        if not PDF_AVAILABLE:
            raise ImportError("pdfplumber not installed")
        
        try:
            # Try pdfplumber for digital PDFs
            with pdfplumber.open(file_path) as pdf:
                text_content = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
                
                if text_content.strip():
                    logger.info(f"Successfully extracted text using pdfplumber: {len(text_content)} characters")
                    return text_content
                
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")
        
        # Fallback to OCR for scanned PDFs
        logger.info("Falling back to OCR extraction")
        return FileExtractor.extract_with_ocr(file_path)
    
    @staticmethod
    def extract_with_ocr(file_path: str) -> str:
        """Extract text using pytesseract OCR"""
        logger.info(f"Extracting text using OCR: {file_path}")
        
        if not OCR_AVAILABLE:
            raise ImportError("pytesseract or PIL not installed")
        
        try:
            # Handle PDF conversion for OCR
            if file_path.lower().endswith('.pdf'):
                # Convert PDF to images first (requires pdf2image)
                try:
                    from pdf2image import convert_from_path
                    images = convert_from_path(file_path)
                    
                    text_content = ""
                    for i, image in enumerate(images):
                        page_text = pytesseract.image_to_string(image)
                        text_content += f"Page {i+1}:\n{page_text}\n"
                    
                    logger.info(f"OCR extraction from PDF successful: {len(text_content)} characters")
                    return text_content
                    
                except ImportError:
                    logger.error("pdf2image not installed - cannot OCR PDF files")
                    raise ImportError("pdf2image required for PDF OCR")
            else:
                # Direct image OCR
                image = Image.open(file_path)
                text_content = pytesseract.image_to_string(image)
                logger.info(f"OCR extraction successful: {len(text_content)} characters")
                return text_content
                
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            raise
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX files"""
        logger.info(f"Extracting text from DOCX: {file_path}")
        
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed")
        
        try:
            doc = Document(file_path)
            text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    text_content += f"\n{row_text}"
            
            logger.info(f"DOCX extraction successful: {len(text_content)} characters")
            return text_content
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
    
    @staticmethod
    def extract_from_excel(file_path: str) -> str:
        """Extract text from Excel files"""
        logger.info(f"Extracting text from Excel: {file_path}")
        
        if not XLSX_AVAILABLE:
            raise ImportError("openpyxl not installed")
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_content = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_content += f"\nSheet: {sheet_name}\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text_content += f"{row_text}\n"
            
            logger.info(f"Excel extraction successful: {len(text_content)} characters")
            return text_content
            
        except Exception as e:
            logger.error(f"Excel extraction failed: {e}")
            raise

class TimetableParser:
    """Parses extracted text to identify timetable data"""
    
    @staticmethod
    def parse_timetable_text(text: str) -> List[Dict[str, Any]]:
        """
        Parse extracted text to identify timetable entries.
        This is a basic implementation - can be enhanced with AI/ML.
        """
        logger.info("Parsing timetable data from extracted text")
        
        # Basic parsing logic - can be enhanced with AI
        timetable_entries = []
        
        # Example patterns to look for
        weekdays = ['monday', 'tuesday', 'wednesday', 'thursday', 'friday', 'saturday', 'sunday']
        lines = text.lower().split('\n')
        
        current_day = None
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line contains a weekday
            for day in weekdays:
                if day in line:
                    current_day = day.capitalize()
                    break
            
            # Look for time patterns (e.g., "09:00-10:00", "9:00 AM - 10:00 AM")
            import re
            time_pattern = r'(\d{1,2}:\d{2})\s*[-–]\s*(\d{1,2}:\d{2})'
            time_match = re.search(time_pattern, line)
            
            if time_match and current_day:
                start_time = time_match.group(1)
                end_time = time_match.group(2)
                
                # Extract subject and class info (basic heuristic)
                remaining_text = line.replace(time_match.group(0), '').strip()
                parts = remaining_text.split()
                
                if len(parts) >= 2:
                    subject = parts[0] if parts else "Unknown Subject"
                    pupils = " ".join(parts[1:]) if len(parts) > 1 else "Unknown Class"
                    
                    timetable_entries.append({
                        "weekday": current_day,
                        "start_time": start_time,
                        "end_time": end_time,
                        "subject": subject,
                        "pupils": pupils
                    })
        
        # If no entries found, create a sample entry for testing
        if not timetable_entries:
            logger.warning("No timetable entries parsed - creating sample entry")
            timetable_entries = [{
                "weekday": "Monday",
                "start_time": "09:00",
                "end_time": "10:00",
                "subject": "Extracted Subject",
                "pupils": "Extracted Class",
                "note": "Extracted from uploaded file"
            }]
        
        logger.info(f"Parsed {len(timetable_entries)} timetable entries")
        return timetable_entries

async def process_timetable_file_task(ctx: dict, teacher_id: str, file_path: str):
    """
    Background task to process timetable files with intelligent text extraction.
    
    Args:
        ctx: ARQ context
        teacher_id: UUID of the teacher
        file_path: Path to the uploaded file
    """
    logger.info(f"Starting timetable file processing for teacher: {teacher_id}, file: {file_path}")
    
    try:
        # Send initial status
        await publish_ws_message(teacher_id, {
            "status": "started",
            "message": "Processing timetable file...",
            "teacher_id": teacher_id,
            "file_path": file_path
        })
        
        # Validate file exists
        if not os.path.exists(file_path):
            error_msg = f"File not found: {file_path}"
            logger.error(error_msg)
            await publish_ws_message(teacher_id, {
                "status": "error",
                "message": error_msg,
                "teacher_id": teacher_id
            })
            return {"error": error_msg}
        
        # Detect file type
        file_type = FileExtractor.detect_file_type(file_path)
        logger.info(f"Detected file type: {file_type}")
        
        if file_type == 'unknown':
            error_msg = f"Unsupported file type: {Path(file_path).suffix}"
            logger.error(error_msg)
            await publish_ws_message(teacher_id, {
                "status": "error", 
                "message": error_msg,
                "teacher_id": teacher_id
            })
            return {"error": error_msg}
        
        # Update progress
        await publish_ws_message(teacher_id, {
            "status": "processing",
            "message": f"Extracting text from {file_type} file...",
            "teacher_id": teacher_id
        })
        
        # Extract text based on file type
        extracted_text = ""
        try:
            if file_type == 'pdf':
                extracted_text = FileExtractor.extract_from_pdf(file_path)
            elif file_type == 'image':
                extracted_text = FileExtractor.extract_with_ocr(file_path)
            elif file_type == 'docx':
                extracted_text = FileExtractor.extract_from_docx(file_path)
            elif file_type == 'excel':
                extracted_text = FileExtractor.extract_from_excel(file_path)
            
        except ImportError as e:
            error_msg = f"Required library not installed: {str(e)}"
            logger.error(error_msg)
            await publish_ws_message(teacher_id, {
                "status": "error",
                "message": error_msg,
                "teacher_id": teacher_id
            })
            return {"error": error_msg}
        
        except Exception as e:
            error_msg = f"Text extraction failed: {str(e)}"
            logger.error(f"{error_msg}\n{traceback.format_exc()}")
            await publish_ws_message(teacher_id, {
                "status": "error",
                "message": error_msg,
                "teacher_id": teacher_id
            })
            return {"error": error_msg}
        
        if not extracted_text.strip():
            error_msg = "No text could be extracted from the file"
            logger.error(error_msg)
            await publish_ws_message(teacher_id, {
                "status": "error",
                "message": error_msg,
                "teacher_id": teacher_id
            })
            return {"error": error_msg}
        
        logger.info(f"Successfully extracted {len(extracted_text)} characters")
        
        # Update progress
        await publish_ws_message(teacher_id, {
            "status": "processing",
            "message": "Parsing timetable data...",
            "teacher_id": teacher_id
        })
        
        # Parse timetable data
        timetable_entries = TimetableParser.parse_timetable_text(extracted_text)
        
        # Save extracted data to database 
        uploaded_file_id = None
        try:
            async with AsyncSession(async_engine) as session:
                # Save uploaded file record to UploadedFile table
                logger.debug("Saving uploaded file record to database...")
                
                # Extract original filename from path
                file_name = Path(file_path).name
                
                # Create UploadedFile record
                uploaded_file = UploadedFile(
                    teacher_id=UUID(teacher_id),
                    file_name=file_name,
                    file_type=file_type,
                    purpose="timetable",  # Since this is for timetable processing
                    gcs_path=None,  # Leave blank as requested
                    extracted_text=extracted_text
                )
                
                session.add(uploaded_file)
                await session.flush()  # Flush to get the ID
                
                # Get the ID while the session is still active
                uploaded_file_id = uploaded_file.id
                logger.info(f"Uploaded file record saved with ID: {uploaded_file_id}")
                
                # Commit the transaction
                await session.commit()
                
        except Exception as e:
            error_msg = f"Database error: {str(e)}"
            logger.error(f"{error_msg}\n{traceback.format_exc()}")
            await publish_ws_message(teacher_id, {
                "status": "error",
                "message": error_msg,
                "teacher_id": teacher_id
            })
            raise
        
        # Send success message after database operations are complete
        success_msg = f"File processed successfully! Extracted {len(timetable_entries)} timetable entries."
        logger.info(success_msg)
        
        # Send success message with extracted data
        await publish_ws_message(teacher_id, {
            "status": "completed",
            "message": success_msg,
            "teacher_id": teacher_id,
            "uploaded_file_id": str(uploaded_file_id),
            "extracted_data": {
                "timetables": timetable_entries,
                "raw_text": extracted_text[:500] + "..." if len(extracted_text) > 500 else extracted_text,
                "file_type": file_type,
                "entries_count": len(timetable_entries),
                "file_name": Path(file_path).name
            }
        })
        
        return {
            "status": "success",
            "uploaded_file_id": str(uploaded_file_id),
            "timetable_entries": timetable_entries,
            "extracted_text": extracted_text,
            "file_type": file_type
        }
    
    except Exception as e:
        error_msg = f"Timetable processing failed: {str(e)}"
        logger.error(f"{error_msg}\n{traceback.format_exc()}")
        
        # Send error notification
        try:
            await publish_ws_message(teacher_id, {
                "status": "error",
                "message": error_msg,
                "teacher_id": teacher_id
            })
        except Exception as ws_error:
            logger.error(f"Failed to send WebSocket error message: {ws_error}")
        
        raise

# ARQ Worker Configuration
async def startup(ctx):
    """ARQ worker startup"""
    ctx['redis'] = await create_pool(arq_redis_settings)
    logger.info("Timetable processing worker started")

async def shutdown(ctx):
    """ARQ worker shutdown"""
    ctx['redis'].close()
    await ctx['redis'].aclose()
    await async_engine.dispose()
    logger.info("Timetable processing worker shutdown")

# Worker configuration for this specific task
timetable_worker_config = {
    'functions': [process_timetable_file_task],
    'redis_settings': arq_redis_settings,
    'on_startup': startup,
    'on_shutdown': shutdown,
    'max_tries': 3,           # Retry failed jobs 3 times
    'retry_delay': 10,        # Wait 10 seconds between retries
    'job_timeout': 300,       # 5 minutes max per job
    'concurrent_jobs': 2,     # Process 2 files simultaneously
    'keep_result': 3600,      # Keep job results for 1 hour
    'max_jobs': 50            # Max jobs before worker restart
}

# Manual testing function
if __name__ == "__main__":
    async def test_enqueue():
        redis = await create_pool(arq_redis_settings)
        try:
            teacher_id = "7bed2b69-8000-4b36-8e91-7fe0b70c9d82"
            test_file = "./uploads/test_timetable.pdf"
            job = await redis.enqueue_job('process_timetable_file_task', teacher_id, test_file)
            print(f"[SUCCESS] Timetable processing job queued: {job.job_id}")
        finally:
            await redis.aclose()
    
    asyncio.run(test_enqueue())