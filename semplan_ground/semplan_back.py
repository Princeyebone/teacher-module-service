"""Background Task for Semester Plan File Processing

This module provides background task processing for semester plan file uploads,
including intelligent text extraction based on file type.

Supported file types:
- PDF: pdfplumber → pytesseract fallback
- Images (JPG/PNG): pytesseract OCR
- DOCX: python-docx
- XLSX: openpyxl
- TXT: plain text

Usage:
    from semplan_back import process_semplan_file_task
    job_id = await enqueue_semplan_processing(teacher_id, file_path, gcs_file_name)
"""

import os
import json
import asyncio
import logging
import traceback
import hashlib
import re
from pathlib import Path
from typing import Dict, List, Optional, Any
from datetime import datetime
from uuid import UUID


# Initialize logger
logger = logging.getLogger(__name__)

# File processing libraries
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    
try:
    import pytesseract
    from PIL import Image
    
    # Configure Tesseract path for Windows
    import platform
    if platform.system() == "Windows":
        # Common Windows installation paths for Tesseract
        possible_paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            r"C:\Users\{username}\AppData\Local\Tesseract-OCR\tesseract.exe".format(
                username=os.environ.get('USERNAME', '')
            ),
            r"C:\Tesseract-OCR\tesseract.exe"
        ]
        
        # Try to find Tesseract executable
        tesseract_path = None
        for path in possible_paths:
            if os.path.exists(path):
                tesseract_path = path
                break
        
        # Set the path if found
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
            logger.info(f"Tesseract found at: {tesseract_path}")
        else:
            logger.warning("Tesseract not found in common paths. Please ensure it's installed and in PATH.")
    
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

# ARQ and database imports
from arq import create_pool, ArqRedis
from arq.connections import RedisSettings
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession
from sqlalchemy import select
import redis.asyncio as redis

# Project imports
from config import settings
from model import TeacherNotification, TempExtract
from sch_ground.background import arq_redis_settings, async_engine, publish_ws_message, save_notification
from database import AsyncSessionLocal


# Logger already initialized at the top

# Initialize async Redis client
redis_client = redis.Redis(host="localhost", port=6379, decode_responses=True)

# File type mappings
SUPPORTED_EXTENSIONS = {
    'pdf': 'pdf',
    'jpg': 'image', 'jpeg': 'image', 'png': 'image', 'bmp': 'image', 'tiff': 'image',
    'docx': 'docx',
    'xlsx': 'excel', 'xls': 'excel',
    'txt': 'text'
}

class FileExtractor:
    """Handles text extraction from various file types"""
    
    @staticmethod
    def detect_file_type(file_path: str) -> str:
        """Detect file type from extension"""
        extension = Path(file_path).suffix.lower().lstrip('.')
        return SUPPORTED_EXTENSIONS.get(extension, 'unknown')
    
    @staticmethod
    def extract_from_text(file_path: str) -> str:
        """Extract text from plain text files"""
        logger.info(f"Extracting text from plain text file: {file_path}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            logger.info(f"✅ Text extraction successful: {len(text_content)} characters")
            return text_content
        except Exception as e:
            logger.error(f"❌ Text extraction failed: {e}")
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text_content = f.read()
                logger.info(f"✅ Text extraction with latin-1 successful: {len(text_content)} characters")
                return text_content
            except Exception as e2:
                logger.error(f"❌ Text extraction with latin-1 also failed: {e2}")
                raise
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF using pdfplumber with pytesseract fallback"""
        logger.info(f"📄 Extracting text from PDF: {file_path}")
        
        if not PDF_AVAILABLE:
            raise ImportError("pdfplumber not installed")
        
        try:
            # Try pdfplumber for digital PDFs
            with pdfplumber.open(file_path) as pdf:
                text_content = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
                
                if text_content.strip():
                    logger.info(f"✅ Successfully extracted text using pdfplumber: {len(text_content)} characters")
                    return text_content
                
        except Exception as e:
            logger.warning(f"⚠️ pdfplumber extraction failed: {e}")
        
        # Fallback to OCR for scanned PDFs
        logger.info("🔄 Falling back to OCR extraction")
        return FileExtractor.extract_with_ocr(file_path)
    
    @staticmethod
    def extract_with_ocr(file_path: str) -> str:
        """Extract text using pytesseract OCR"""
        logger.info(f"🔍 Extracting text using OCR: {file_path}")
        
        if not OCR_AVAILABLE:
            raise ImportError("pytesseract or PIL not installed")
        
        try:
            # Handle PDF conversion for OCR
            if file_path.lower().endswith('.pdf'):
                # Convert PDF to images first (requires pdf2image)
                try:
                    from pdf2image import convert_from_path
                    images = convert_from_path(file_path)
                    
                    text_content = ""
                    for i, image in enumerate(images):
                        page_text = pytesseract.image_to_string(image)
                        text_content += f"Page {i+1}:\n{page_text}\n"
                    
                    logger.info(f"✅ OCR extraction from PDF successful: {len(text_content)} characters")
                    return text_content
                    
                except ImportError:
                    logger.error("❌ pdf2image not installed - cannot OCR PDF files")
                    raise ImportError("pdf2image required for PDF OCR")
            else:
                # Direct OCR for image files
                image = Image.open(file_path)
                text_content = pytesseract.image_to_string(image)
                logger.info(f"✅ OCR extraction successful: {len(text_content)} characters")
                return text_content
                
        except Exception as e:
            logger.error(f"❌ OCR extraction failed: {e}\n{traceback.format_exc()}")
            raise
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX files"""
        logger.info(f"📝 Extracting text from DOCX: {file_path}")
        
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed")
        
        try:
            doc = Document(file_path)
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + "\t"
                    text_content += "\n"
            
            logger.info(f"✅ DOCX extraction successful: {len(text_content)} characters")
            return text_content
            
        except Exception as e:
            logger.error(f"❌ DOCX extraction failed: {e}\n{traceback.format_exc()}")
            raise
    
    @staticmethod
    def extract_from_excel(file_path: str) -> str:
        """Extract text from Excel files"""
        logger.info(f"📊 Extracting text from Excel: {file_path}")
        
        if not XLSX_AVAILABLE:
            raise ImportError("openpyxl not installed")
        
        try:
            workbook = openpyxl.load_workbook(file_path, read_only=True)
            text_content = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_content += f"Sheet: {sheet_name}\n"
                
                for row in sheet.iter_rows(values_only=True):
                    for cell_value in row:
                        if cell_value is not None:
                            text_content += str(cell_value) + "\t"
                    text_content += "\n"
                text_content += "\n"
            
            logger.info(f"✅ Excel extraction successful: {len(text_content)} characters")
            return text_content
            
        except Exception as e:
            logger.error(f"❌ Excel extraction failed: {e}\n{traceback.format_exc()}")
            raise

async def process_semplan_file_task(ctx: Dict[Any, Any], teacher_id: str, file_path: str, gcs_file_name: str, subject: str = None, class_name: str = None, session_data: Dict = None) -> Dict[str, Any]:
    """
    Background task to process semester plan file and extract text.
    
    Args:
        ctx: ARQ context
        teacher_id: Teacher UUID string
        file_path: Local path to the file
        gcs_file_name: Name of the file in GCS
        subject: Subject name (optional)
        class_name: Class name (optional)
        session_data: Semester dates and other data (optional)
    
    Returns:
        Dict with processing results
    """
    logger.info(f"[SEMPLAN] Starting processing task for teacher {teacher_id}")
    logger.info(f"[SEMPLAN] File: {file_path}, GCS Name: {gcs_file_name}")
    logger.info(f"[SEMPLAN] Subject: {subject}, Class: {class_name}")
    logger.info(f"[SEMPLAN] Session data: {session_data}")
    
    # Global set to track sent messages across all function instances
    GLOBAL_SENT_MESSAGES = set()
    
    # Track sent messages to prevent duplicates
    sent_messages = set()
    
    async def send_unique_ws_message(teacher_id: str, message: dict):
        """Send WebSocket message only if it hasn't been sent before"""
        # Create a more robust hash of the message content to identify duplicates
        # Include more fields to make the hash more unique
        status = message.get('status', '')
        type_ = message.get('type', '')
        msg_text = message.get('message', '')
        subject_ = str(message.get('subject', ''))
        class_name_ = str(message.get('class_name', ''))
        entry_count = str(message.get('extracted_characters', ''))  # Convert to string for consistency
        
        message_key = f"{status}_{type_}_{msg_text}_{subject_}_{class_name_}_{entry_count}"
        
        # Use a more robust deduplication approach
        message_hash = hashlib.md5(message_key.encode()).hexdigest()
        
        # Check both local and global deduplication
        logger.info(f"🔍 [DEDUPLICATION] Checking message hash: {message_hash} for key: {message_key}")
        logger.info(f"🔍 [DEDUPLICATION] Local sent_messages set size: {len(sent_messages)}")
        logger.info(f"🔍 [DEDUPLICATION] Global SENT_MESSAGES set size: {len(GLOBAL_SENT_MESSAGES)}")
        
        if message_hash not in sent_messages and message_hash not in GLOBAL_SENT_MESSAGES:
            sent_messages.add(message_hash)
            GLOBAL_SENT_MESSAGES.add(message_hash)
            logger.info(f"➕ [DEDUPLICATION] Adding message hash to both local and global sent_messages: {message_hash}")
            logger.info(f"➕ [DEDUPLICATION] Local set now contains {len(sent_messages)} items")
            logger.info(f"➕ [DEDUPLICATION] Global set now contains {len(GLOBAL_SENT_MESSAGES)} items")
            try:
                logger.info(f"Sending WebSocket message to teacher {teacher_id}")
                logger.info(f"Message content: {message}")
                await publish_ws_message(teacher_id, message)
                logger.info(f"✅ [DEDUPLICATION] Sent unique WebSocket message with hash: {message_hash}")
            except Exception as e:
                logger.error(f"❌ [DEDUPLICATION] Failed to send WebSocket message: {e}")
                logger.error(f"Full traceback: {traceback.format_exc()}")
            return True
        else:
            logger.info(f"⏭️ [DEDUPLICATION] SKIPPING DUPLICATE MESSAGE with hash: {message_hash}")
            logger.info(f"⏭️ [DEDUPLICATION] Message already sent, not sending again")
            if message_hash in sent_messages:
                logger.info(f"⏭️ [DEDUPLICATION] Message found in local sent_messages")
            if message_hash in GLOBAL_SENT_MESSAGES:
                logger.info(f"⏭️ [DEDUPLICATION] Message found in global SENT_MESSAGES")
            return False
    
    try:
        # Notify start of processing
        await send_unique_ws_message(
            teacher_id,
            {
                "type": "semplan_processing",
                "status": "started",
                "message": "Starting semester plan text extraction...",
                "file_name": os.path.basename(file_path),
                "teacher_id": teacher_id,
                "subject": subject,
                "class_name": class_name
            }
        )
        
        # Extract session data if subject and class_name are provided
        enhanced_session_data = session_data or {}
        if subject and class_name:
            logger.info(f"[SEMPLAN] Extracting session data for subject: {subject}, class: {class_name}")
            
            # Create database session
            async with AsyncSession(async_engine) as db_session:
                try:
                    # Get academic calendar data
                    from model import AcademicCalendar
                    from sqlalchemy import select
                    
                    acc = (await db_session.execute(
                        select(AcademicCalendar).where(AcademicCalendar.teacher_id == UUID(teacher_id))
                    )).scalar_one_or_none()
                    
                    if acc:
                        semester_start_date = acc.semester_start_date
                        semester_end_date = acc.semester_end_date
                        logger.info(f"[SEMPLAN] Academic calendar: start_date={semester_start_date}, end_date={semester_end_date}")
                        
                        # Get class sessions for the specified subject and class
                        from model import ClassSession
                        class_sessions = (await db_session.execute(
                            select(ClassSession).where(
                                (ClassSession.subject.ilike(f"%{subject}%")) &
                                (ClassSession.class_name.ilike(f"%{class_name}%")) &
                                (ClassSession.teacher_id == UUID(teacher_id))
                            )
                        )).scalars().all()
                        
                        logger.info(f"[SEMPLAN] Found {len(class_sessions)} class sessions for {subject} - {class_name}")
                        
                        # Log details of each session for debugging
                        for i, session_obj in enumerate(class_sessions):
                            logger.info(f"[SEMPLAN] Session {i+1}: ID={session_obj.id}, Date={session_obj.date}, "
                                      f"Start={session_obj.start_time}, End={session_obj.end_time}, "
                                      f"Subject={session_obj.subject}, Class={session_obj.class_name}")
                        
                        # Group sessions by week
                        sessions_by_week = {}
                        
                        for session_obj in class_sessions:
                            # Calculate week number based on academic calendar start date
                            # Using timedelta to ensure accurate week calculation
                            days_diff = (session_obj.date - semester_start_date).days
                            week_number = (days_diff // 7) + 1
                            
                            # Ensure week number is within valid range (1-11 or 16)
                            if 1 <= week_number <= 16:
                                week_key = f"Week {week_number}"
                                
                                if week_key not in sessions_by_week:
                                    sessions_by_week[week_key] = []
                                
                                # Create session info object with date information
                                session_info = {
                                    "id": session_obj.id,
                                    "date": str(session_obj.date),
                                    "subject": session_obj.subject,
                                    "start_time": str(session_obj.start_time),
                                    "end_time": str(session_obj.end_time),
                                    "class_name": session_obj.class_name,
                                    "location": session_obj.location,
                                    "session_number": session_obj.session_number
                                }
                                
                                sessions_by_week[week_key].append(session_info)
                                logger.info(f"[SEMPLAN] Assigned session {session_obj.id} (date: {session_obj.date}) to {week_key}")
                            else:
                                logger.warning(f"[SEMPLAN] Session {session_obj.id} (date: {session_obj.date}) falls outside valid week range (1-11 or 16), week calculated as {week_number}")
                        
                        logger.info(f"[SEMPLAN] Grouped sessions into {len(sessions_by_week)} weeks")
                        
                        # Log detailed structure of sessions_by_week for debugging
                        logger.info(f"[SEMPLAN] Sessions by week structure:")
                        for week_key, sessions in sessions_by_week.items():
                            logger.info(f"[SEMPLAN]   {week_key}: {len(sessions)} sessions")
                            for i, session in enumerate(sessions):
                                logger.info(f"[SEMPLAN]     Session {i+1}: ID={session['id']}, Date={session['date']}, "
                                          f"Subject={session['subject']}, Class={session['class_name']}, "
                                          f"Start={session['start_time']}, End={session['end_time']}, "
                                          f"Location={session['location']}, Session#={session['session_number']}")
                        
                        # Prepare enhanced session data for background task
                        # Filter session data to only include essential fields for AI
                        weekly_session_data = {}
                        for week_key, sessions in sessions_by_week.items():
                            # Create filtered sessions with only essential fields
                            filtered_sessions = [
                                {
                                    "id": session["id"],
                                    "date": session["date"],
                                    "start_time": session["start_time"],
                                    "end_time": session["end_time"],
                                    "week_number": int(week_key.replace("Week ", ""))
                                }
                                for session in sessions
                            ]
                            
                            weekly_session_data[week_key] = {
                                "week_number": int(week_key.replace("Week ", "")),
                                "sessions": filtered_sessions
                            }
                        
                        # Log detailed structure of weekly_session_data for debugging
                        logger.info(f"[SEMPLAN] Weekly session data structure being sent to AI:")
                        logger.info(f"[SEMPLAN] COMPLETE WEEKLY SESSION DATA: {json.dumps(weekly_session_data, indent=2, default=str)}")
                        for week_key, week_data in weekly_session_data.items():
                            logger.info(f"[SEMPLAN]   {week_key}: Week#={week_data['week_number']}, {len(week_data['sessions'])} sessions")
                            for i, session in enumerate(week_data['sessions']):
                                logger.info(f"[SEMPLAN]     Session {i+1}: ID={session['id']}, Date={session['date']}, "
                                          f"Start={session['start_time']}, End={session['end_time']}, "
                                          f"Week#={session['week_number']}")
                        
                        # Update session_data with weekly sessions
                        enhanced_session_data.update({
                            "semester_start_date": str(semester_start_date),
                            "semester_end_date": str(semester_end_date),
                            "weekly_sessions": weekly_session_data
                        })
                        
                        logger.info(f"[SEMPLAN] Enhanced session data with {len(weekly_session_data)} weeks of sessions")
                        # Log summary of sessions per week
                        for week_key, sessions in weekly_session_data.items():
                            logger.info(f"[SEMPLAN] Week {sessions['week_number']}: {len(sessions['sessions'])} sessions")
                    else:
                        logger.warning(f"[SEMPLAN] No academic calendar found for teacher {teacher_id}")
                except Exception as e:
                    logger.error(f"[SEMPLAN] Error extracting session data: {e}")
                    logger.error(f"[SEMPLAN] Full traceback: {traceback.format_exc()}")
        
        # Detect file type
        file_type = FileExtractor.detect_file_type(file_path)
        logger.info(f"[SEMPLAN] Detected file type: {file_type}")
        
        if file_type == 'text':
            extracted_text = FileExtractor.extract_from_text(file_path)
        elif file_type == 'pdf':
            extracted_text = FileExtractor.extract_from_pdf(file_path)
        elif file_type == 'image':
            extracted_text = FileExtractor.extract_with_ocr(file_path)
        elif file_type == 'docx':
            extracted_text = FileExtractor.extract_from_docx(file_path)
        elif file_type == 'excel':
            extracted_text = FileExtractor.extract_from_excel(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Log extraction results
        logger.info(f"[SEMPLAN] Text extraction completed. Extracted {len(extracted_text)} characters")
        # Log a sample of the extracted text for debugging (first 500 characters)
        logger.info(f"[SEMPLAN] SAMPLE OF EXTRACTED TEXT: {extracted_text[:500]}...")
        # Note: Even if extracted_text is unclear or incomplete, the AI will also analyze the original GCS file
        if len(extracted_text) < 100:
            logger.warning(f"[SEMPLAN] Extracted text appears to be incomplete or unclear ({len(extracted_text)} characters). AI will also analyze the original file at GCS for complete information.")
        
        # DELETE THE FILE AFTER EXTRACTION - as per requirements
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"[SEMPLAN] Deleted temporary file: {file_path}")
            else:
                logger.warning(f"[SEMPLAN] Temporary file not found for deletion: {file_path}")
        except Exception as e:
            logger.error(f"[SEMPLAN] Failed to delete temporary file {file_path}: {e}")
        
        # Process with AI if we have necessary parameters (extracted text is not required if we have the GCS file)
        ai_result = None
        # Changed condition: Process with AI if we have subject and class_name, regardless of extracted_text clarity
        if subject and class_name:
            try:
                from external_service import send_semester_plan_to_ai
                if hasattr(settings, 'API_KEY') and settings.API_KEY:
                    logger.info("🤖 Sending data to AI for semester plan processing (will use both extracted text and GCS file)")
                    # Log session data being sent to AI for debugging
                    if session_data:
                        logger.info(f"📅 SESSION DATA BEING SENT TO AI:")
                        logger.info(f"   Semester Start: {session_data.get('semester_start_date', 'Not provided')}")
                        logger.info(f"   Semester End: {session_data.get('semester_end_date', 'Not provided')}")
                        weekly_sessions = session_data.get('weekly_sessions', {})
                        logger.info(f"   Number of Weeks: {len(weekly_sessions)}")
                        for week_key, week_data in weekly_sessions.items():
                            logger.info(f"     {week_key}: {len(week_data.get('sessions', []))} sessions")
                        # Log the actual weeks available to make it clear to the AI
                        available_weeks = list(weekly_sessions.keys())
                        logger.info(f"   🔑 AVAILABLE WEEKS FOR MAPPING (AI MUST USE ONLY THESE): {', '.join(available_weeks)}")
                    else:
                        logger.warning("⚠️ NO SESSION DATA PROVIDED TO AI")
                    
                    # Validate that we have weekly session data
                    weekly_sessions = enhanced_session_data.get('weekly_sessions', {})
                    if not weekly_sessions:
                        logger.warning("⚠️ NO WEEKLY SESSION DATA FOUND - AI may not be able to map sessions correctly")
                    else:
                        logger.info(f"✅ Found {len(weekly_sessions)} weeks of session data to send to AI")
                        # Log which weeks we have
                        week_numbers = list(weekly_sessions.keys())
                        logger.info(f"📅 Weeks available: {week_numbers}")
                        logger.info("📌 IMPORTANT: These are the ONLY weeks the AI should use for mapping, regardless of weeks mentioned in the document")
                        
                        # Validate that session data contains the required fields
                        for week_key, week_data in weekly_sessions.items():
                            sessions = week_data.get('sessions', [])
                            if sessions:
                                # Check first session for required fields
                                first_session = sessions[0]
                                required_fields = ['id', 'date', 'start_time', 'end_time', 'week_number']
                                missing_fields = [field for field in required_fields if field not in first_session]
                                if missing_fields:
                                    logger.warning(f"⚠️ Missing fields in session data for {week_key}: {missing_fields}")
                                else:
                                    logger.info(f"✅ Session data for {week_key} contains all required fields")
                            else:
                                logger.info(f"📅 No sessions found for {week_key}")
                    
                    # Send to AI with both extracted text (if available) and GCS file path
                    # NOTE: The AI has been instructed to ONLY use the weekly_sessions data for mapping,
                    # and to IGNORE any week numbers mentioned in the extracted text or document
                    ai_result = send_semester_plan_to_ai(
                        extracted_text,  # This may be empty or unclear, but AI will also use the GCS file
                        f"gs://{settings.GCS_BUCKET_NAME}/{gcs_file_name}",
                        settings.API_KEY,
                        enhanced_session_data,
                        class_name,
                        subject
                    )
                    
                    if "error" not in ai_result:
                        logger.info("🎉 AI processing successful")
                        # Log the complete AI response for debugging
                        logger.info(f"🤖 COMPLETE AI RESPONSE: {json.dumps(ai_result, indent=2, default=str)}")
                        # Store AI response directly in the Strand/Substrand/ContentStandard/Indicator tables
                        try:
                            await store_ai_response_in_tables(
                                teacher_id, 
                                class_name, 
                                subject, 
                                ai_result
                            )
                            # Notify completion only if AI processing and storage were successful
                            await send_unique_ws_message(
                                teacher_id,
                                {
                                    "type": "semplan_processing",
                                    "status": "completed",
                                    "message": f"Text extraction completed successfully. Extracted {len(extracted_text)} characters.",
                                    "file_name": os.path.basename(file_path),
                                    "teacher_id": teacher_id,
                                    "subject": subject,
                                    "class_name": class_name,
                                    "extracted_characters": len(extracted_text),
                                    "ai_processed": True
                                }
                            )
                        except Exception as storage_error:
                            logger.error(f"💥 Failed to store AI response in tables: {storage_error}")
                            logger.error(f"Full traceback: {traceback.format_exc()}")
                            # Notify error
                            await send_unique_ws_message(
                                teacher_id,
                                {
                                    "type": "semplan_processing",
                                    "status": "error",
                                    "message": f"Failed to store AI response in tables: {str(storage_error)}",
                                    "file_name": os.path.basename(file_path),
                                    "teacher_id": teacher_id,
                                    "subject": subject,
                                    "class_name": class_name
                                }
                            )
                            # Re-raise the exception to be caught by the outer exception handler
                            raise
                    else:
                        logger.warning(f"⚠️ AI processing failed: {ai_result['error']}")
                        # Log the error response for debugging
                        logger.warning(f"🤖 AI ERROR RESPONSE: {json.dumps(ai_result, indent=2, default=str)}")
                        # Notify AI processing error
                        await send_unique_ws_message(
                            teacher_id,
                            {
                                "type": "semplan_processing",
                                "status": "error",
                                "message": f"AI processing failed: {ai_result['error']}",
                                "file_name": os.path.basename(file_path),
                                "teacher_id": teacher_id,
                                "subject": subject,
                                "class_name": class_name
                            }
                        )
                else:
                    logger.info("⏭️ Skipping AI processing - no API key configured")
                    # Notify completion when skipping AI processing
                    await send_unique_ws_message(
                        teacher_id,
                        {
                            "type": "semplan_processing",
                            "status": "completed",
                            "message": f"Text extraction completed successfully. Extracted {len(extracted_text)} characters. AI processing skipped (no API key).",
                            "file_name": os.path.basename(file_path),
                            "teacher_id": teacher_id,
                            "subject": subject,
                            "class_name": class_name,
                            "extracted_characters": len(extracted_text),
                            "ai_processed": False
                        }
                    )
            except Exception as e:
                logger.error(f"💥 AI processing failed: {e}")
                logger.error(f"Full traceback: {traceback.format_exc()}")
                # Notify error
                await send_unique_ws_message(
                    teacher_id,
                    {
                        "type": "semplan_processing",
                        "status": "error",
                        "message": f"AI processing failed: {str(e)}",
                        "file_name": os.path.basename(file_path),
                        "teacher_id": teacher_id,
                        "subject": subject,
                        "class_name": class_name
                    }
                )
        else:
            logger.info("⏭️ Skipping AI processing - missing subject or class_name")
            # Notify completion when skipping AI processing
            await send_unique_ws_message(
                teacher_id,
                {
                    "type": "semplan_processing",
                    "status": "completed",
                    "message": f"Text extraction completed successfully. Extracted {len(extracted_text)} characters. AI processing skipped (missing subject or class).",
                    "file_name": os.path.basename(file_path),
                    "teacher_id": teacher_id,
                    "subject": subject,
                    "class_name": class_name,
                    "extracted_characters": len(extracted_text),
                    "ai_processed": False
                }
            )
        
        # Return results WITHOUT the extracted text (as per requirements)
        # The extracted text will be provided to a prompt later, not stored anywhere
        return {
            "status": "success",
            "teacher_id": teacher_id,
            "file_path": file_path,  # This file no longer exists
            "gcs_file_name": gcs_file_name,
            "file_type": file_type,
            "extracted_characters": len(extracted_text),  # Only return character count
            "subject": subject,
            "class_name": class_name,
            "session_data": enhanced_session_data,  # Include enhanced session data for prompt builder
            "ai_result": ai_result,  # Include AI result if processing was successful
            "message": "Text extraction completed successfully. File deleted after extraction."
        }
        
    except Exception as e:
        error_msg = f"Semester plan processing failed: {str(e)}"
        logger.error(f"[SEMPLAN] {error_msg}\n{traceback.format_exc()}")
        
        # DELETE THE FILE EVEN IF PROCESSING FAILED - as per requirements
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"[SEMPLAN] Deleted temporary file after error: {file_path}")
            else:
                logger.warning(f"[SEMPLAN] Temporary file not found for deletion after error: {file_path}")
        except Exception as delete_error:
            logger.error(f"[SEMPLAN] Failed to delete temporary file after error {file_path}: {delete_error}")
        
        # Notify error
        await send_unique_ws_message(
            teacher_id,
            {
                "type": "semplan_processing",
                "status": "error",
                "message": error_msg,
                "file_name": os.path.basename(file_path) if 'file_path' in locals() else "Unknown",
                "teacher_id": teacher_id,
                "subject": subject,
                "class_name": class_name
            }
        )
        
        raise Exception(error_msg)
    finally:
        # Clean up global sent messages for this function instance
        try:
            for message_hash in sent_messages:
                if message_hash in GLOBAL_SENT_MESSAGES:
                    GLOBAL_SENT_MESSAGES.remove(message_hash)
            logger.info(f"🧹 Cleaned up {len(sent_messages)} message hashes from global set")
        except Exception as e:
            logger.warning(f"⚠️ Error cleaning up global sent messages: {e}")

# Create separate Redis settings (without queue_name parameter)
from arq.connections import RedisSettings
semplan_redis_settings = RedisSettings(host="localhost", port=6379, database=0, conn_timeout=10, conn_retries=5, conn_retry_delay=1)

# ARQ Worker Configuration
async def startup(ctx):
    """ARQ worker startup"""
    # Create pool with specific queue name
    ctx['redis'] = await create_pool(semplan_redis_settings, default_queue_name='semplan_queue')
    logger.info("Semester plan processing worker started")

async def shutdown(ctx):
    """ARQ worker shutdown"""
    ctx['redis'].close()
    await ctx['redis'].aclose()
    await async_engine.dispose()
    logger.info("Semester plan processing worker shutdown")

semplan_worker_config = {
    "functions": [process_semplan_file_task],
    "redis_settings": semplan_redis_settings,
    "queue_name": "semplan_queue",  # Use just the queue name without arq:queue: prefix
    "on_startup": startup,
    "on_shutdown": shutdown,
    "max_tries": 3,
    "job_timeout": 300,  # 5 minutes
    "concurrent_jobs": 5,
    "retry_delay": 60,   # 1 minute
    "keep_result": 3600, # 1 hour
    "poll_delay": 0.5,
}

# Function to enqueue semester plan processing task
async def enqueue_semplan_processing(teacher_id: str, file_path: str, gcs_file_name: str, subject: str = None, class_name: str = None, session_data: Dict = None) -> Optional[str]:
    """
    Enqueue a semester plan processing task.
    
    Args:
        teacher_id: Teacher UUID string
        file_path: Local path to the file
        gcs_file_name: Name of the file in GCS
        subject: Subject name (optional)
        class_name: Class name (optional)
        session_data: Weekly session data (optional)
    
    Returns:
        Job ID if successful, None if failed
    """
    logger.info(f"[ENQUEUE] Enqueuing semester plan processing for teacher {teacher_id}")
    logger.info(f"[ENQUEUE] File: {file_path}, GCS Name: {gcs_file_name}")
    logger.info(f"[ENQUEUE] Subject: {subject}, Class: {class_name}")
    
    try:
        # Create ARQ redis pool
        redis_pool = await create_pool(semplan_redis_settings)
        
        # Enqueue the task with all parameters and specify the queue name
        job = await redis_pool.enqueue_job(
            'process_semplan_file_task',
            teacher_id,
            file_path,
            gcs_file_name,
            subject,
            class_name,
            session_data,
            _queue_name='semplan_queue'
        )
        
        if job:
            logger.info(f"[ENQUEUE] Successfully enqueued job: {job.job_id}")
            return job.job_id
        else:
            logger.error("[ENQUEUE] Failed to enqueue job - job creation returned None")
            return None
            
    except Exception as e:
        logger.error(f"[ENQUEUE] Failed to enqueue semester plan processing: {e}\n{traceback.format_exc()}")
        return None

# Function to store AI response in TempExtract table
async def store_ai_response_in_temp_extract(teacher_id: str, class_name: str, subject: str, ai_response: Dict, gcs_file_name: str = None) -> None:
    """
    Store AI response in TempExtract table with type 'semester plan'.
    Ensures only one instance exists for a teacher, class, and subject combination.
    
    Args:
        teacher_id: Teacher UUID string
        class_name: Class name
        subject: Subject name
        ai_response: AI response data to store
        gcs_file_name: GCS file name for generating signed URL (optional)
    """
    logger.info(f"[SEMPLAN] Storing AI response in TempExtract for teacher {teacher_id}, class {class_name}, subject {subject}")
    
    try:
        # Create database session
        async with AsyncSession(async_engine) as session:
            # Check if an entry already exists for this teacher, class, and subject with type 'semester plan'
            from sqlalchemy import select
            from model import TempExtract
            from uuid import UUID
            
            existing_entry = (await session.execute(
                select(TempExtract).where(
                    (TempExtract.teacher_id == UUID(teacher_id)) &
                    (TempExtract.class_name == class_name) &
                    (TempExtract.subject == subject) &
                    (TempExtract.type == "semester plan")
                )
            )).scalar_one_or_none()
            
            # Generate signed URL if gcs_file_name is provided
            signed_url = None
            if gcs_file_name:
                try:
                    from gcs_utils import generate_signed_url
                    from config import settings
                    # Generate a signed URL that expires in 7 days (604800 seconds)
                    # For GET requests, we don't need to include special headers
                    signed_url = generate_signed_url(
                        settings.GCS_BUCKET_NAME, 
                        gcs_file_name, 
                        method="GET",
                        expiration=604800,
                        only_include_host_in_headers=False
                    )
                    logger.info(f"[SEMPLAN] Generated signed URL for file: {gcs_file_name}")
                except Exception as e:
                    logger.error(f"[SEMPLAN] Failed to generate signed URL: {e}")
                    signed_url = None
            
            if existing_entry:
                # Update existing entry
                logger.info(f"[SEMPLAN] Updating existing TempExtract entry: {existing_entry.id}")
                existing_entry.data = ai_response
                existing_entry.file = signed_url  # Store the signed URL
                existing_entry.updated_at = datetime.utcnow()
                session.add(existing_entry)
            else:
                # Create new entry
                logger.info(f"[SEMPLAN] Creating new TempExtract entry")
                new_entry = TempExtract(
                    teacher_id=UUID(teacher_id),
                    type="semester plan",
                    class_name=class_name,
                    subject=subject,
                    file=signed_url,  # Store the signed URL
                    data=ai_response
                )
                session.add(new_entry)
            
            # Commit the transaction
            await session.commit()
            logger.info(f"[SEMPLAN] AI response stored successfully in TempExtract")
            
    except Exception as e:
        logger.error(f"[SEMPLAN] Failed to store AI response in TempExtract: {e}")
        logger.error(f"[SEMPLAN] Full traceback: {traceback.format_exc()}")
        raise

# New function to store AI response directly in the Strand/Substrand/ContentStandard/Indicator tables
async def store_ai_response_in_tables(teacher_id: str, class_name: str, subject: str, ai_response: dict):
    """
    Store the AI response data directly in the Strand, Substrand, ContentStandard, and Indicator tables.
    
    Args:
        teacher_id: The UUID of the teacher
        class_name: The class name
        subject: The subject name
        ai_response: The AI response containing strand_data, substrand_data, content_standard_data, and indicator_data
    """
    try:
        logger.info(f"[SEMPLAN] Storing AI response directly in tables for teacher {teacher_id}, class {class_name}, subject {subject}")
        logger.info(f"[SEMPLAN] COMPLETE AI RESPONSE BEING STORED: {json.dumps(ai_response, indent=2, default=str)}")
        
        # Create a new database session
        async with AsyncSessionLocal() as db_session:
            # First, delete any existing data for this teacher, class, and subject combination
            # This ensures we have a clean slate for the new AI-generated data
            from sqlalchemy import and_, delete
            from model import Strand, Substrand, ContentStandard, Indicator
            
            # Delete in the correct order to avoid foreign key constraint violations
            # Delete indicators first (no foreign key dependencies)
            await db_session.execute(
                delete(Indicator).where(
                    and_(
                        Indicator.teacher_id == UUID(teacher_id),
                        Indicator.class_name == class_name,
                        Indicator.subject == subject
                    )
                )
            )
            # Then delete content standards
            await db_session.execute(
                delete(ContentStandard).where(
                    and_(
                        ContentStandard.teacher_id == UUID(teacher_id),
                        ContentStandard.class_name == class_name,
                        ContentStandard.subject == subject
                    )
                )
            )
            # Then delete substrands
            await db_session.execute(
                delete(Substrand).where(
                    and_(
                        Substrand.teacher_id == UUID(teacher_id),
                        Substrand.class_name == class_name,
                        Substrand.subject == subject
                    )
                )
            )
            # Finally delete strands
            await db_session.execute(
                delete(Strand).where(
                    and_(
                        Strand.teacher_id == UUID(teacher_id),
                        Strand.class_name == class_name,
                        Strand.subject == subject
                    )
                )
            )
            await db_session.commit()
            logger.info(f"[SEMPLAN] Deleted existing data for teacher {teacher_id}, class {class_name}, subject {subject}")
            
            # Process strand data - handle both flat and nested structures
            strand_data_list = ai_response.get("strand_data", [])
            logger.info(f"[SEMPLAN] Processing {len(strand_data_list)} strand entries")
            created_strands = {}  # strand_name -> {week_number -> strand_object}
            
            # Flatten nested strand data if needed
            flattened_strand_data = []
            for strand_data in strand_data_list:
                # Check if this is nested data (has substrand_data inside)
                if "substrand_data" in strand_data:
                    # This is nested data, extract the strand info
                    flattened_strand_data.append({
                        "id": strand_data.get("id", ""),
                        "strand_name": strand_data.get("name", strand_data.get("strand_name", "")),
                        "name": strand_data.get("name", ""),  # Keep for backward compatibility
                        "weeks": strand_data.get("weeks", []),
                        "session_ids": strand_data.get("session_ids", []),
                        "session_details": strand_data.get("session_details", [])
                    })
                    
                    # Extract substrand, content standard, and indicator data from nested structure
                    for substrand_data in strand_data.get("substrand_data", []):
                        # Add substrand data to the flat list
                        substrand_entry = {
                            "id": substrand_data.get("id", ""),
                            "strand_name": strand_data.get("name", strand_data.get("strand_name", "")),
                            "substrand_name": substrand_data.get("name", substrand_data.get("substrand_name", "")),
                            "name": substrand_data.get("name", ""),  # Keep for backward compatibility
                            "weeks": substrand_data.get("weeks", []),
                            "session_ids": substrand_data.get("session_ids", []),
                            "session_details": substrand_data.get("session_details", [])
                        }
                        
                        # Extract content standards and indicators from nested structure
                        for cs_data in substrand_data.get("content_standard_data", []):
                            # Add content standard data to the flat list
                            cs_entry = {
                                "id": cs_data.get("id", ""),
                                "strand_name": strand_data.get("name", strand_data.get("strand_name", "")),
                                "substrand_name": substrand_data.get("name", substrand_data.get("substrand_name", "")),
                                "content_standard_code": cs_data.get("id", cs_data.get("content_standard_code", "")),
                                "content_standard_text": cs_data.get("name", cs_data.get("content_standard_text", "")),
                                "name": cs_data.get("name", ""),  # Keep for backward compatibility
                                "session_ids": cs_data.get("session_ids", []),
                                "session_details": cs_data.get("session_details", [])
                            }
                            
                            # Extract indicators from nested structure
                            for indicator_data in cs_data.get("indicator_data", []):
                                # Add indicator data to the flat list
                                indicator_entry = {
                                    "id": indicator_data.get("id", ""),
                                    "strand_name": strand_data.get("name", strand_data.get("strand_name", "")),
                                    "substrand_name": substrand_data.get("name", substrand_data.get("substrand_name", "")),
                                    "content_standard_code": cs_data.get("id", cs_data.get("content_standard_code", "")),
                                    "indicator_code": indicator_data.get("id", indicator_data.get("indicator_code", "")),
                                    "indicator_text": indicator_data.get("name", indicator_data.get("indicator_text", "")),
                                    "name": indicator_data.get("name", ""),  # Keep for backward compatibility
                                    "session_ids": indicator_data.get("session_ids", []),
                                    "session_details": indicator_data.get("session_details", [])
                                }
                                
                                # Extract session information for indicators
                                session_ids = []
                                session_details = []
                                for session_item in indicator_data.get("sessions", []):
                                    session_ids.append(session_item.get("id", 0))
                                    session_details.append({
                                        "id": session_item.get("id", 0),
                                        "date": session_item.get("date", ""),
                                        "start_time": session_item.get("start_time", ""),
                                        "end_time": session_item.get("end_time", ""),
                                        "week_number": session_item.get("week_number", 0)
                                    })
                                
                                indicator_entry["session_ids"] = session_ids
                                indicator_entry["session_details"] = session_details
                                
                                # Add to the flat indicator data list (will be processed later)
                                if "indicator_data" not in ai_response:
                                    ai_response["indicator_data"] = []
                                ai_response["indicator_data"].append(indicator_entry)
                            
                            # Extract session information for content standards
                            session_ids = []
                            session_details = []
                            for session in cs_data.get("sessions", []):
                                session_ids.append(session.get("id", 0))
                                session_details.append({
                                    "id": session.get("id", 0),
                                    "date": session.get("date", ""),
                                    "start_time": session.get("start_time", ""),
                                    "end_time": session.get("end_time", ""),
                                    "week_number": session.get("week_number", 0)
                                })
                            
                            cs_entry["session_ids"] = session_ids
                            cs_entry["session_details"] = session_details
                            
                            # Add to the flat content standard data list (will be processed later)
                            if "content_standard_data" not in ai_response:
                                ai_response["content_standard_data"] = []
                            ai_response["content_standard_data"].append(cs_entry)
                        
                        # Extract session information for substrands
                        session_ids = []
                        session_details = []
                        for session in substrand_data.get("sessions", []):
                            session_ids.append(session.get("id", 0))
                            session_details.append({
                                "id": session.get("id", 0),
                                "date": session.get("date", ""),
                                "start_time": session.get("start_time", ""),
                                "end_time": session.get("end_time", ""),
                                "week_number": session.get("week_number", 0)
                            })
                        
                        substrand_entry["session_ids"] = session_ids
                        substrand_entry["session_details"] = session_details
                        
                        # Add to the flat substrand data list (will be processed later)
                        if "substrand_data" not in ai_response:
                            ai_response["substrand_data"] = []
                        ai_response["substrand_data"].append(substrand_entry)
                    
                    # Extract session information for strands
                    session_ids = []
                    session_details = []
                    for session in strand_data.get("sessions", []):
                        session_ids.append(session.get("id", 0))
                        session_details.append({
                            "id": session.get("id", 0),
                            "date": session.get("date", ""),
                            "start_time": session.get("start_time", ""),
                            "end_time": session.get("end_time", ""),
                            "week_number": session.get("week_number", 0)
                        })
                    
                    # Update the flattened strand entry with session info
                    if flattened_strand_data:
                        flattened_strand_data[-1]["session_ids"] = session_ids
                        flattened_strand_data[-1]["session_details"] = session_details
                else:
                    # This is already flat data
                    flattened_strand_data.append(strand_data)
            
            # Use the flattened strand data
            strand_data_list = flattened_strand_data
            
            for strand_data in strand_data_list:
                logger.info(f"[SEMPLAN] Processing strand data: {json.dumps(strand_data, indent=2, default=str)}")
                # Handle both "strand_name" and "name" fields for backward compatibility
                strand_name = strand_data.get("strand_name", strand_data.get("name", ""))
                weeks = strand_data.get("weeks", [])
                session_ids = strand_data.get("session_ids", [])
                session_details = strand_data.get("session_details", [])
                
                # If no sessions were provided directly, try to extract them from the sessions data
                if not session_ids and "sessions" in strand_data:
                    session_ids = []
                    session_details = []
                    for session in strand_data.get("sessions", []):
                        session_ids.append(session.get("id", 0))
                        session_details.append({
                            "id": session.get("id", 0),
                            "date": session.get("date", ""),
                            "start_time": session.get("start_time", ""),
                            "end_time": session.get("end_time", ""),
                            "week_number": session.get("week_number", 0)
                        })
                    logger.info(f"[SEMPLAN] Extracted {len(session_ids)} sessions from strand data")
                
                # Convert weeks to integers for storage - handle both numeric and text formats
                week_numbers_int = []
                for week in weeks:
                    try:
                        # If it's already an integer, use it directly
                        if isinstance(week, int):
                            week_numbers_int.append(week)
                        # If it's a string, try to extract the numeric part
                        elif isinstance(week, str):
                            # Handle formats like "Week 3", "week3", "3", etc.
                            numeric_match = re.search(r'\d+', week)
                            if numeric_match:
                                week_numbers_int.append(int(numeric_match.group()))
                            else:
                                # If no numeric part found, try to convert directly
                                week_numbers_int.append(int(week))
                        # If it's something else, try to convert to int
                        else:
                            week_numbers_int.append(int(week))
                    except (ValueError, TypeError) as e:
                        logger.warning(f"[SEMPLAN] Invalid week format in strand '{strand_name}': {week} - {str(e)}")
                
                # Log the converted week numbers
                logger.info(f"[SEMPLAN] Strand '{strand_name}' uses weeks: {week_numbers_int}")
                
                # VALIDATION: Ensure strand has a valid name
                if not strand_name or not strand_name.strip():
                    logger.warning(f"[SEMPLAN] Skipping strand with empty name")
                    continue
                
                # VALIDATION: Ensure strand has at least one week
                if not week_numbers_int:
                    logger.warning(f"[SEMPLAN] Strand '{strand_name}' has no valid weeks assigned")
                    # Skip this strand as it's not valid
                    continue
                
                # If no sessions were provided directly, try to extract them from the weeks data
                if not session_ids:
                    logger.info(f"[SEMPLAN] No direct session IDs for strand '{strand_name}', checking for session data in weeks")
                
                # Create strand entries - one for each week
                strand_entries = []
                for week_num in week_numbers_int:
                    strand_entry = Strand(
                        strand_name=strand_name,
                        subject=subject,
                        class_name=class_name,
                        teacher_id=UUID(teacher_id),
                        week_number=week_num,  # Single week number for each strand entry
                        session_ids=session_ids,
                        session_details=session_details
                    )
                    db_session.add(strand_entry)
                    await db_session.flush()  # Get the ID
                    strand_entries.append(strand_entry)
                    logger.debug(f"[SEMPLAN] Created strand entry with ID {strand_entry.id}")
                
                # Store in our tracking dictionary
                if strand_name not in created_strands:
                    created_strands[strand_name] = {}
                for i, week_num in enumerate(week_numbers_int):
                    created_strands[strand_name][week_num] = strand_entries[i]
            
            # Process substrand data - handle both flat and nested structures
            substrand_data_list = ai_response.get("substrand_data", [])
            logger.info(f"[SEMPLAN] Processing {len(substrand_data_list)} substrand entries")
            created_substrands = {}  # substrand_name -> substrand_object
            
            for substrand_data in substrand_data_list:
                logger.info(f"[SEMPLAN] Processing substrand data: {json.dumps(substrand_data, indent=2, default=str)}")
                # Handle both "strand_name" and "name" fields for backward compatibility
                strand_name = substrand_data.get("strand_name", substrand_data.get("strand_name", ""))
                # Handle both "substrand_name" and "name" fields for backward compatibility
                substrand_name = substrand_data.get("substrand_name", substrand_data.get("name", ""))
                weeks = substrand_data.get("weeks", [])
                session_ids = substrand_data.get("session_ids", [])
                session_details = substrand_data.get("session_details", [])
                
                # If no sessions were provided directly, try to extract them from the sessions data
                if not session_ids and "sessions" in substrand_data:
                    session_ids = []
                    session_details = []
                    for session in substrand_data.get("sessions", []):
                        session_ids.append(session.get("id", 0))
                        session_details.append({
                            "id": session.get("id", 0),
                            "date": session.get("date", ""),
                            "start_time": session.get("start_time", ""),
                            "end_time": session.get("end_time", ""),
                            "week_number": session.get("week_number", 0)
                        })
                    logger.info(f"[SEMPLAN] Extracted {len(session_ids)} sessions from substrand data")
                
                # Convert weeks to integers for storage - handle both numeric and text formats
                week_numbers_int = []
                for week in weeks:
                    try:
                        # If it's already an integer, use it directly
                        if isinstance(week, int):
                            week_numbers_int.append(week)
                        # If it's a string, try to extract the numeric part
                        elif isinstance(week, str):
                            # Handle formats like "Week 3", "week3", "3", etc.
                            numeric_match = re.search(r'\d+', week)
                            if numeric_match:
                                week_numbers_int.append(int(numeric_match.group()))
                            else:
                                # If no numeric part found, try to convert directly
                                week_numbers_int.append(int(week))
                        # If it's something else, try to convert to int
                        else:
                            week_numbers_int.append(int(week))
                    except (ValueError, TypeError) as e:
                        logger.warning(f"[SEMPLAN] Invalid week format in substrand '{substrand_name}': {week} - {str(e)}")
                
                # Log the converted week numbers
                logger.info(f"[SEMPLAN] Substrand '{substrand_name}' uses weeks: {week_numbers_int}")
                
                # VALIDATION: Ensure substrand has a valid name
                if not substrand_name or not substrand_name.strip():
                    logger.warning(f"[SEMPLAN] Skipping substrand with empty name")
                    continue
                
                # VALIDATION: Ensure substrand has at least one week
                if not week_numbers_int:
                    logger.warning(f"[SEMPLAN] Substrand '{substrand_name}' has no valid weeks assigned")
                    # Skip this substrand as it's not valid
                    continue
                
                # If no sessions were provided directly, try to extract them from the weeks data
                if not session_ids:
                    logger.info(f"[SEMPLAN] No direct session IDs for substrand '{substrand_name}', checking for session data in weeks")
                
                # VALIDATION: Ensure substrand weeks are within its parent strand's weeks
                strand_weeks = []
                if strand_name in created_strands:
                    strand_weeks = list(created_strands[strand_name].keys())
                
                invalid_weeks = [week for week in week_numbers_int if week not in strand_weeks]
                if invalid_weeks:
                    logger.warning(f"[SEMPLAN] Substrand '{substrand_name}' has weeks not in parent strand: {invalid_weeks}")
                
                # Find the strand to link the substrand
                strand_id = 0
                if strand_name in created_strands and weeks:
                    # Use the first valid week to find the strand
                    first_valid_week = weeks[0]
                    try:
                        if isinstance(first_valid_week, int):
                            first_valid_week = first_valid_week
                        elif isinstance(first_valid_week, str):
                            numeric_match = re.search(r'\d+', first_valid_week)
                            if numeric_match:
                                first_valid_week = int(numeric_match.group())
                            else:
                                first_valid_week = int(first_valid_week)
                        else:
                            first_valid_week = int(first_valid_week)
                    except (ValueError, TypeError):
                        logger.warning(f"[SEMPLAN] Invalid week format for substrand {substrand_name}: {first_valid_week}")
                    
                    if first_valid_week in created_strands[strand_name]:
                        strand_id = created_strands[strand_name][first_valid_week].id
                elif strand_name in created_strands:
                    # Fallback to first available strand week
                    strand_weeks = list(created_strands[strand_name].keys())
                    if strand_weeks:
                        strand_id = created_strands[strand_name][strand_weeks[0]].id
                
                if strand_id == 0:
                    logger.warning(f"[SEMPLAN] No strand found to link substrand {substrand_name}")
                    continue
                
                # Create substrand entry
                substrand_entry = Substrand(
                    substrand_name=substrand_name,
                    strand_id=strand_id,
                    subject=subject,
                    class_name=class_name,
                    teacher_id=UUID(teacher_id),
                    week_numbers=week_numbers_int,
                    session_ids=session_ids,
                    session_details=session_details
                )
                db_session.add(substrand_entry)
                await db_session.flush()  # Get the ID
                created_substrands[substrand_name] = substrand_entry
                logger.debug(f"[SEMPLAN] Created substrand entry with ID {substrand_entry.id}")
            
            # Process content standard data
            content_standard_data_list = ai_response.get("content_standard_data", [])
            logger.info(f"[SEMPLAN] Processing {len(content_standard_data_list)} content standard entries")
            created_content_standards = {}  # (substrand_name, code) -> content_standard_object
            
            for cs_data in content_standard_data_list:
                logger.info(f"[SEMPLAN] Processing content standard data: {json.dumps(cs_data, indent=2, default=str)}")
                strand_name = cs_data.get("strand_name", cs_data.get("strand_name", ""))
                substrand_name = cs_data.get("substrand_name", cs_data.get("substrand_name", ""))
                # Handle both "content_standard_code" and "id" fields for backward compatibility
                content_standard_code = cs_data.get("content_standard_code", cs_data.get("id", ""))
                # Handle both "content_standard_text" and "name" fields for backward compatibility
                content_standard_text = cs_data.get("content_standard_text", cs_data.get("name", ""))
                session_ids = cs_data.get("session_ids", [])
                session_details = cs_data.get("session_details", [])
                
                # If no sessions were provided directly, try to extract them from the sessions data
                if not session_ids and "sessions" in cs_data:
                    session_ids = []
                    session_details = []
                    for session in cs_data.get("sessions", []):
                        session_ids.append(session.get("id", 0))
                        session_details.append({
                            "id": session.get("id", 0),
                            "date": session.get("date", ""),
                            "start_time": session.get("start_time", ""),
                            "end_time": session.get("end_time", ""),
                            "week_number": session.get("week_number", 0)
                        })
                    logger.info(f"[SEMPLAN] Extracted {len(session_ids)} sessions from content standard data")
                
                # Log session details for debugging
                logger.info(f"[SEMPLAN] Content standard '{content_standard_code}' session details: {len(session_details)} sessions")
                
                # If no sessions were provided directly, try to extract them from the weeks data
                if not session_ids:
                    logger.info(f"[SEMPLAN] No direct session IDs for content standard '{content_standard_code}', checking for session data in weeks")
                
                # VALIDATION: Ensure content standard has at least one session
                if not session_ids or len(session_ids) == 0:
                    logger.warning(f"[SEMPLAN] Content standard '{content_standard_code}' has no sessions assigned")
                elif len(session_ids) > 0:
                    logger.info(f"[SEMPLAN] Content standard '{content_standard_code}' has {len(session_ids)} sessions assigned")
                
                # VALIDATION: Ensure content standard weeks are within its parent strand's weeks
                # We need to find the substrand to get its strand
                substrand_weeks = []
                if substrand_name in created_substrands:
                    substrand_obj = created_substrands[substrand_name]
                    substrand_weeks = substrand_obj.week_numbers
                
                # Find the strand for validation
                strand_weeks = []
                if strand_name in created_strands:
                    strand_weeks = list(created_strands[strand_name].keys())
                
                # Check if content standard weeks are valid
                # (This is a bit more complex as content standards don't directly store weeks,
                # but we can validate based on their parent substrand/strand)
                
                if not content_standard_text:
                    logger.warning(f"[SEMPLAN] Skipping content standard with no text")
                    continue
                
                # Find the substrand to link the content standard
                substrand_id = 0
                if substrand_name in created_substrands:
                    substrand_id = created_substrands[substrand_name].id
                
                if substrand_id == 0:
                    logger.warning(f"[SEMPLAN] No substrand found to link content standard {content_standard_code}")
                    continue
                
                # Create content standard entry
                cs_entry = ContentStandard(
                    content_standard_code=content_standard_code or None,
                    content_standard=content_standard_text,
                    substrand_id=substrand_id,
                    subject=subject,
                    class_name=class_name,
                    teacher_id=UUID(teacher_id),
                    session_ids=session_ids,
                    session_details=session_details
                )
                db_session.add(cs_entry)
                await db_session.flush()  # Get the ID
                key = (substrand_name, content_standard_code)
                created_content_standards[key] = cs_entry
                logger.debug(f"[SEMPLAN] Created content standard entry with ID {cs_entry.id}")
            
            # Process indicator data
            indicator_data_list = ai_response.get("indicator_data", [])
            logger.info(f"[SEMPLAN] Processing {len(indicator_data_list)} indicator entries")
            
            # First, let's validate that indicators follow the session assignment rules
            # Group indicators by session ID to check if any session has more than 2 indicators
            session_indicator_count = {}
            for indicator_data in indicator_data_list:
                # Extract session IDs from either direct field or sessions array
                session_ids = indicator_data.get("session_ids", [])
                # If no direct session_ids, try to extract from sessions array
                if not session_ids and "sessions" in indicator_data:
                    session_ids = [session.get("id", 0) for session in indicator_data.get("sessions", [])]
                
                for session_id in session_ids:
                    if session_id not in session_indicator_count:
                        session_indicator_count[session_id] = 0
                    session_indicator_count[session_id] += 1
            
            # Log sessions with more than 2 indicators
            for session_id, count in session_indicator_count.items():
                if count > 2:
                    logger.warning(f"[SEMPLAN] Session {session_id} has {count} indicators assigned (maximum recommended is 2)")
                elif count == 0:
                    logger.warning(f"[SEMPLAN] Indicator with session_id {session_id} has no sessions assigned")
                else:
                    logger.info(f"[SEMPLAN] Session {session_id} has {count} indicators assigned")
            
            for indicator_data in indicator_data_list:
                logger.info(f"[SEMPLAN] Processing indicator data: {json.dumps(indicator_data, indent=2, default=str)}")
                strand_name = indicator_data.get("strand_name", indicator_data.get("strand_name", ""))
                substrand_name = indicator_data.get("substrand_name", indicator_data.get("substrand_name", ""))
                content_standard_code = indicator_data.get("content_standard_code", indicator_data.get("content_standard_code", ""))
                # Handle both "indicator_code" and "id" fields for backward compatibility
                indicator_code = indicator_data.get("indicator_code", indicator_data.get("id", ""))
                # Handle both "indicator_text" and "name" fields for backward compatibility
                indicator_text = indicator_data.get("indicator_text", indicator_data.get("name", ""))
                # Extract session data from either direct fields or sessions array
                session_ids = indicator_data.get("session_ids", [])
                session_details = indicator_data.get("session_details", [])
                
                # If no sessions were provided directly, try to extract them from the sessions data
                if not session_ids and "sessions" in indicator_data:
                    session_ids = []
                    session_details = []
                    for session in indicator_data.get("sessions", []):
                        session_ids.append(session.get("id", 0))
                        session_details.append({
                            "id": session.get("id", 0),
                            "date": session.get("date", ""),
                            "start_time": session.get("start_time", ""),
                            "end_time": session.get("end_time", ""),
                            "week_number": session.get("week_number", 0)
                        })
                    logger.info(f"[SEMPLAN] Extracted {len(session_ids)} sessions from indicator data")
                
                # Log session details for debugging
                logger.info(f"[SEMPLAN] Indicator '{indicator_code}' session details: {len(session_details)} sessions")
                
                # VALIDATION: Ensure indicator has at least one session
                if not session_ids or len(session_ids) == 0:
                    logger.warning(f"[SEMPLAN] Indicator '{indicator_code}' has no sessions assigned")
                elif len(session_ids) > 2:
                    logger.warning(f"[SEMPLAN] Indicator '{indicator_code}' has {len(session_ids)} sessions assigned (maximum recommended is 2)")
                elif len(session_ids) == 1:
                    logger.info(f"[SEMPLAN] Indicator '{indicator_code}' has {len(session_ids)} session assigned (ideal)")
                elif len(session_ids) == 2:
                    logger.info(f"[SEMPLAN] Indicator '{indicator_code}' has {len(session_ids)} sessions assigned (paired with another indicator)")
                
                # VALIDATION: Ensure indicator weeks are within its parent strand's weeks
                # (Similar to content standards, indicators don't directly store weeks)
                
                if not indicator_text:
                    logger.warning(f"[SEMPLAN] Skipping indicator with no text")
                    continue
                
                # Find the content standard to link the indicator
                content_standard_id = 0
                key = (substrand_name, content_standard_code)
                if key in created_content_standards:
                    content_standard_id = created_content_standards[key].id
                
                if content_standard_id == 0:
                    logger.warning(f"[SEMPLAN] No content standard found to link indicator {indicator_code}")
                    continue
                
                # Create indicator entry
                indicator_entry = Indicator(
                    indicator_code=indicator_code or None,
                    indicator_text=indicator_text,
                    content_standard_id=content_standard_id,
                    subject=subject,
                    class_name=class_name,
                    teacher_id=UUID(teacher_id),
                    session_ids=session_ids,
                    session_details=session_details
                )
                db_session.add(indicator_entry)
                logger.debug(f"[SEMPLAN] Created indicator entry with ID {indicator_entry.id}")
            
            # Commit all changes
            await db_session.commit()
            logger.info(f"[SEMPLAN] AI response stored successfully in tables")
            
    except Exception as e:
        logger.error(f"[SEMPLAN] Failed to store AI response in tables: {e}")
        logger.error(f"[SEMPLAN] Full traceback: {traceback.format_exc()}")
        raise